"""Render a :class:`FactFindContext` into the bundled Fact-Find ``.docx``.

The template is filled by *label matching*, not absolute coordinates: we walk the
document body in order, track the current product-section heading and the last
sub-heading paragraph, and dispatch each table to a handler chosen from its own
header text. This survives minor template edits (added rows, reordered cells)
far better than hard-coded indices.

Cells we cannot resolve are left exactly as the template has them — blank, for
the broker to complete by hand.
"""

from __future__ import annotations

import copy
import io
import re
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from sqlalchemy.orm import Session

from app.models import PolicyYear
from app.services.fact_find_form import (
    AGE_BANDS,
    MAX_BASIS_ROWS,
    PAGE_SECTIONS,
    TEMPLATE_PATH,
    build_context,
    section_for_code,
)

if TYPE_CHECKING:
    from app.services.fact_find_form import FactFindContext, SectionContext

_NUMERAL_RE = re.compile(r"^([ivx]+)\)", re.I)

# Lowercase roman numerals for renumbering expanded basis rows (i) to xx)).
_ROMAN = (
    "i",
    "ii",
    "iii",
    "iv",
    "v",
    "vi",
    "vii",
    "viii",
    "ix",
    "x",
    "xi",
    "xii",
    "xiii",
    "xiv",
    "xv",
    "xvi",
    "xvii",
    "xviii",
    "xix",
    "xx",
)


def _roman(n: int) -> str:
    return _ROMAN[n - 1] if 1 <= n <= len(_ROMAN) else str(n)


# Paragraph-heading text → section code. Matched as a prefix (case-insensitive).
_HEADINGS: tuple[tuple[str, str], ...] = (
    ("GENERAL INFORMATION", "GI"),
    ("GROUP TERM LIFE", "GTL"),
    ("GROUP PERSONAL ACCIDENT", "GPA"),
    ("GROUP HOSPITAL & SURGICAL", "GHS"),
    ("GROUP CATASTROPHIC MEDICAL", "GCM"),
    ("GROUP CLINICAL GENERAL PRACTITIONER", "GCGP_GCSP"),
    ("GROUP BUSINESS TRAVEL", "GBT"),
    ("NEEDS ANALYSIS", "_NEEDS"),
    ("DECLARATION", "_DECL"),
)


def _heading_section(text: str) -> str | None:
    up = text.strip().upper()
    for prefix, code in _HEADINGS:
        if up.startswith(prefix):
            return code
    return None


def _ctext(cell: _Cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


def _set(cell: _Cell, text: str | int) -> None:
    """Write ``text`` into a (normally empty) cell, keeping its paragraph style."""
    s = str(text)
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = s
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(s)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _mark(cell: _Cell) -> None:
    _set(cell, "X")


def _set_compact(cell: _Cell, text: str) -> None:
    """Write a monetary value into the template's narrow amount field."""
    _set(cell, text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(7)
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(cell._tc.makeelement(qn("w:noWrap"), {}))


def _keep_table_together(table: Table) -> None:
    """Keep compact tick-range tables from splitting between Word pages."""
    for row in table.rows[:-1]:
        for cell in _physical_cells(row):
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def _physical_cells(row) -> list[_Cell]:
    """Cells as they exist visually, without python-docx merge aliases.

    ``row.cells`` repeats the same cell object for every grid column covered by
    a merge. Writing by those repeated indices overwrote labels such as
    ``from`` and ``Employees only`` instead of their adjacent blank fields.
    """
    cells: list[_Cell] = []
    seen: set[int] = set()
    for cell in row.cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        cells.append(cell)
    return cells


def _remove_unconfigured_pages(doc, ctx: FactFindContext) -> None:
    """Remove product-page blocks that are not configured for this year."""
    section = "GI"
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag.rsplit("}", 1)[-1] == "p":
            heading = _heading_section(Paragraph(child, doc).text.strip())
            if heading is not None:
                section = heading
        if section in PAGE_SECTIONS and section not in ctx.sections:
            body.remove(child)


def render_docx(ctx: FactFindContext) -> bytes:
    doc = Document(str(TEMPLATE_PATH))
    _remove_unconfigured_pages(doc, ctx)

    section = "GI"
    last_para = ""
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            txt = Paragraph(child, doc).text.strip()
            if txt:
                hs = _heading_section(txt)
                if hs is not None:
                    section = hs
                last_para = txt
        elif tag == "tbl":
            table = Table(child, doc)
            _fill_table(ctx, section, last_para, table)

    _normalise_pagination(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _header_text(table: Table) -> str:
    if not table.rows:
        return ""
    return " | ".join(_ctext(c) for c in table.rows[0].cells)


def _fill_table(ctx: FactFindContext, section: str, last_para: str, table: Table) -> None:
    header = _header_text(table)
    first = _ctext(table.rows[0].cells[0]) if table.rows else ""

    if section in {"GI", "_NEEDS"}:
        if first.startswith("Name of Company"):
            _fill_company(ctx, table)
        elif first.startswith("Insurance Coverage"):
            _fill_matrix(ctx, table)
        elif first.startswith("Company") and ctx.other_products:
            _fill_other_products(ctx, table)
        return

    sec = ctx.sections.get(section)
    if sec is None:
        return  # product not configured on this policy year → leave page blank

    if first.startswith("Presently Insured"):
        _fill_presently_insured(sec, table)
    elif first.startswith("Eligibility"):
        _fill_eligibility(sec, table)
    elif first.startswith("Date employees can join scheme"):
        _fill_entry_date(sec, table)
    elif "up to age" in header.lower() and sec.code == "GTL":
        _fill_free_cover_limit(sec, table)
    elif "Age Band" in header:
        _fill_age_band(sec, table)
    elif "Highest Sum Insured" in header:
        _fill_highest_sum(sec, table)
    elif "oldest insured member" in header:
        _fill_oldest_sum(sec, table)
    elif "Singaporeans" in header:
        _fill_family(sec, table, sec.family_local)
    elif "Holders of EP" in header or "EP, SP" in header:
        _fill_family(sec, table, sec.family_foreign)
    elif first == "No. of Insured Members":
        _fill_family(sec, table, _combined_family(sec))
    elif first.startswith("Category of Employees") and last_para.strip().lower() == (
        "basis of cover"
    ):
        _fill_basis(sec, table)
    elif first.startswith("Category of Employees") and "Sum Insured" in header:
        _fill_basis(sec, table)
    elif first.startswith("Category of Employees") and sec.code == "GBT":
        _fill_travel_basis(sec, table, last_para)
    elif first.startswith("Period of Coverage"):
        _fill_claim_summaries(sec, table)
    elif first.startswith("Date of Claim"):
        _fill_claim_details(sec, table)


# ── General Information ──────────────────────────────────────────────────────
def _fill_company(ctx: FactFindContext, table: Table) -> None:
    for row in table.rows:
        cells = _physical_cells(row)
        label = _ctext(cells[0]) if cells else ""
        if label.startswith("Name of Company") and len(cells) > 1:
            _set(cells[1], ctx.company_name)
        elif label.startswith("Company Address") and len(cells) > 1:
            _set(cells[1], ctx.company_address)
        elif label.startswith("Nature of Business") and len(cells) > 1:
            _set(cells[1], ctx.nature_of_business)
        elif label.startswith("Country of Origin"):
            if len(cells) > 1 and ctx.country_of_origin:
                _set(cells[1], ctx.country_of_origin)
            if len(cells) > 3:
                _set(cells[3], ctx.total_employees)


def _matrix_row_text(product) -> str:
    return f"{product.title} ({product.code})\n-  for employees\n-  for dependants"


def _apply_matrix_modes(employee_row, dependant_row, product) -> None:
    employee_cells = _physical_cells(employee_row)
    dependant_cells = _physical_cells(dependant_row)
    _set(employee_cells[0], _matrix_row_text(product))
    for cells in (employee_cells, dependant_cells):
        if len(cells) > 1:
            _set(cells[1], "")
        if len(cells) > 2:
            _set(cells[2], "")
    for cells, mode in (
        (employee_cells, product.employee_participation),
        (dependant_cells, product.dependant_participation),
    ):
        if mode:
            _mark(cells[2 if mode.startswith("vol") else 1])


def _split_clinical_matrix_rows(ctx: FactFindContext, table: Table) -> None:
    """Give GCGP and GCSP independent employee/dependant row pairs.

    Their detailed benefits intentionally share one legacy form section, but
    the products and plan participation are configured independently.  The
    template's combined matrix line therefore cannot represent the source data.
    """
    products = [
        ctx.matrix_products[code]
        for code in ("GCGP", "GCSP")
        if code in ctx.matrix_products
    ]
    if not products:
        return
    rows = table.rows
    start = next(
        (
            index
            for index, row in enumerate(rows[:-1])
            if "(GCGP & GCSP)" in _ctext(row.cells[0])
        ),
        None,
    )
    if start is None:
        return
    _apply_matrix_modes(rows[start], rows[start + 1], products[0])
    insert_after = rows[start + 1]._tr
    for product in products[1:]:
        employee_xml = copy.deepcopy(rows[start]._tr)
        dependant_xml = copy.deepcopy(rows[start + 1]._tr)
        insert_after.addnext(employee_xml)
        employee_xml.addnext(dependant_xml)
        insert_after = dependant_xml
        refreshed = table.rows
        _apply_matrix_modes(refreshed[start + 2], refreshed[start + 3], product)
        start += 2


def _fill_matrix(ctx: FactFindContext, table: Table) -> None:
    """Mark Compulsory/Voluntary for each configured product line.

    Each product occupies two physical rows: the first (carrying the 'Group …'
    label) is the employees row, the row immediately after is dependants.
    """
    _split_clinical_matrix_rows(ctx, table)
    rows = table.rows
    i = 0
    while i < len(rows):
        label = _ctext(rows[i].cells[0])
        m = re.search(r"\(([A-Z][A-Z &]*)\)", label)
        raw_code = _canon_matrix_code(m.group(1)) if m else ""
        product = ctx.matrix_products.get(raw_code)
        if product and i + 1 < len(rows):
            _apply_matrix_modes(rows[i], rows[i + 1], product)
            i += 2
            continue
        code = section_for_code(raw_code) if raw_code else None
        if code and code in ctx.sections:
            sec = ctx.sections[code]
            employee_cells = _physical_cells(rows[i])
            mode = sec.employee_participation or sec.participation
            col = 2 if (mode or "").startswith("vol") else 1
            if mode and len(employee_cells) > col:
                _mark(employee_cells[col])
            if sec.has_dependants and i + 1 < len(rows):
                dependant_cells = _physical_cells(rows[i + 1])
                dep_mode = sec.dependant_participation or mode
                dep_col = 2 if (dep_mode or "").startswith("vol") else 1
                if dep_mode and len(dependant_cells) > dep_col:
                    _mark(dependant_cells[dep_col])
            i += 2
        else:
            i += 1

    if ctx.other_products:
        # Extend the matrix with the configured products the legacy form calls
        # "Others". Clone the final employee/dependant pair so borders and row
        # heights remain native to the template.
        source = rows[-2:]
        for product in ctx.other_products:
            clones = [copy.deepcopy(row._tr) for row in source]
            table._tbl.append(clones[0])
            table._tbl.append(clones[1])
            employee_row, dependant_row = table.rows[-2:]
            emp_cells = _physical_cells(employee_row)
            dep_cells = _physical_cells(dependant_row)
            _set(
                emp_cells[0],
                f"{product.title} ({product.code})\n-  for employees\n-  for dependants",
            )
            for cells, mode in (
                (emp_cells, product.employee_participation),
                (dep_cells, product.dependant_participation),
            ):
                # Clear copied marks, then apply this product's opinion.
                if len(cells) > 1:
                    _set(cells[1], "")
                if len(cells) > 2:
                    _set(cells[2], "")
                if mode:
                    _mark(cells[2 if mode.startswith("vol") else 1])


def _fill_other_products(ctx: FactFindContext, table: Table) -> None:
    cells = _physical_cells(table.rows[0]) if table.rows else []
    if not cells:
        return
    last = cells[-1]
    _set(last, "Others:\n" + ", ".join(product.code for product in ctx.other_products))


def _canon_matrix_code(raw: str) -> str:
    """Normalise a matrix label code like 'GCGP & GCSP' to a single code."""
    raw = raw.strip().upper()
    return raw.split("&")[0].strip()


_FORM_TITLE = "EMPLOYEE INSURANCE FACT-FIND FORM"
_LEGACY_PAGE_MARKER_RE = re.compile(r"^INSPRO\s*(?:/|[-\u2013\u2014])", re.I)


def _remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _is_disposable_blank(paragraph: Paragraph) -> bool:
    if paragraph.text.strip():
        return False
    protected = ("drawing", "pict", "object", "fldChar", "instrText", "tab", "br")
    return not any(paragraph._p.xpath(f".//w:{name}") for name in protected)


def _normalise_pagination(doc) -> None:
    """Replace the template's fixed three-page layout with content-aware flow.

    Populated tables grow beyond the legacy form's fixed slots.  Hard internal
    breaks plus dozens of blank positioning paragraphs then create near-empty
    overflow pages and let the following product start midway down a page.
    Section headers are kept with the fields that follow, while all contents
    paginate naturally according to their rendered height.
    """
    for paragraph in list(doc.paragraphs):
        for page_break in paragraph._p.xpath('.//w:br[@w:type="page"]'):
            page_break.getparent().remove(page_break)
        p_pr = paragraph._p.pPr
        if p_pr is not None:
            page_break_before = p_pr.find(qn("w:pageBreakBefore"))
            if page_break_before is not None:
                p_pr.remove(page_break_before)
        text = " ".join(paragraph.text.split())
        if _LEGACY_PAGE_MARKER_RE.match(text):
            _remove_paragraph(paragraph)

    blank_run = 0
    for child in list(doc.element.body.iterchildren()):
        if child.tag.rsplit("}", 1)[-1] != "p":
            blank_run = 0
            continue
        paragraph = Paragraph(child, doc)
        if not _is_disposable_blank(paragraph):
            blank_run = 0
            continue
        blank_run += 1
        if blank_run > 1:
            doc.element.body.remove(child)

    # Do not orphan a repeated form title or DECLARATION heading at the bottom
    # of a page.  Keep its introductory paragraph block with the first table;
    # Word will move the block only when the remaining page space is too small.
    keep_header = False
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "tbl":
            keep_header = False
            continue
        if tag != "p":
            continue
        paragraph = Paragraph(child, doc)
        text = " ".join(paragraph.text.split())
        if text in {_FORM_TITLE, "DECLARATION"}:
            keep_header = True
        if keep_header and text:
            paragraph.paragraph_format.keep_with_next = True


# ── Per-product pages ────────────────────────────────────────────────────────
def _fill_presently_insured(sec: SectionContext, table: Table) -> None:
    for row in table.rows:
        cells = _physical_cells(row)
        label = _ctext(cells[0]) if cells else ""
        if label.startswith("Presently Insured") and len(cells) >= 4:
            # [label][cross][Yes][cross][No]
            _mark(cells[1] if sec.insurer else cells[3])
        elif label.startswith("If Yes") and "insurer" in label and len(cells) > 1:
            _set(cells[1], sec.insurer)
        elif label.startswith("Period of Insurance"):
            # [Period of Insurance:][from][value][to][value]
            if len(cells) >= 5:
                if sec.period_from:
                    _set(cells[2], sec.period_from)
                if sec.period_to:
                    _set(cells[4], sec.period_to)


def _fill_eligibility(sec: SectionContext, table: Table) -> None:
    for row in table.rows:
        cells = _physical_cells(row)
        joined = " ".join(_ctext(c) for c in cells)
        is_emp_only = "Employees only" in joined and "Dependants" not in joined
        is_emp_dep = "Employees and Dependants" in joined
        if not (is_emp_only or is_emp_dep):
            continue
        want = is_emp_dep if sec.has_dependants else is_emp_only
        if not want:
            continue
        # Physical layouts are [cross][label][employee label][value] and, for
        # dependant cover, [cross][label][employee label][value][dep label][value].
        if cells:
            _mark(cells[0])
        if len(cells) > 3:
            _set(cells[3], sec.employees_count)
        if is_emp_dep and len(cells) > 5:
            _set(cells[5], sec.dependants_count)


def _fill_entry_date(sec: SectionContext, table: Table) -> None:
    text = sec.eligibility_date.strip()
    if not text:
        return
    lower = text.lower()
    for row in table.rows:
        cells = _physical_cells(row)
        joined = " ".join(_ctext(cell) for cell in cells)
        if "On date of employment" in joined and (
            "upon employment" in lower or "date of employment" in lower
        ):
            _mark(cells[0])
            return
    # No structured probation value was captured; preserve the platform's
    # exact eligibility wording in the free-text Others line.
    cells = _physical_cells(table.rows[-1])
    if len(cells) >= 3:
        _mark(cells[0])
        _set(cells[2], text)


def _fill_free_cover_limit(sec: SectionContext, table: Table) -> None:
    if not table.rows:
        return
    cells = _physical_cells(table.rows[0])
    if len(cells) >= 3:
        if sec.free_cover_limit:
            _set(cells[0], sec.free_cover_limit)
        if sec.nel_age_limit:
            _set(cells[2], sec.nel_age_limit)


def _is_vmerge_continuation_row(row, anchor_col: int) -> bool:
    """True when ``row`` continues a vertical merge in its anchor column.

    Reads the raw ``<w:tc>`` of the physical row (not python-docx's merge-resolved
    cell, which returns the origin and so always looks like a ``restart``). A
    ``<w:vMerge>`` without ``val="restart"`` marks a continuation row.
    """
    tcs = row._tr.findall(qn("w:tc"))
    if anchor_col >= len(tcs):
        return False
    tc_pr = tcs[anchor_col].find(qn("w:tcPr"))
    if tc_pr is None:
        return False
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is None:
        return False
    return v_merge.get(qn("w:val")) != "restart"


def _basis_units(table: Table, anchor_col: int) -> list[list[int]]:
    """Group physical row indices into category units.

    A unit begins at a numbered origin row and absorbs the vertical-merge
    continuation rows beneath it (the GCGP/GCSP page stacks 3 physical rows per
    category). Header / Total / blank rows are ignored.
    """
    units: list[list[int]] = []
    for i, r in enumerate(table.rows):
        if not _NUMERAL_RE.match(_ctext(r.cells[0])):
            continue
        if _is_vmerge_continuation_row(r, anchor_col):
            if units:
                units[-1].append(i)
        else:
            units.append([i])
    return units


def _clone_last_unit(table: Table, units: list[list[int]]) -> None:
    """Deep-copy the last category unit's ``<w:tr>`` rows and append them after
    it, preserving borders, column widths and the vertical-merge structure."""
    insert_after = table.rows[units[-1][-1]]._tr
    for ridx in units[-1]:
        clone = copy.deepcopy(table.rows[ridx]._tr)
        insert_after.addnext(clone)
        insert_after = clone


def _write_designation(
    cells: list[_Cell],
    c_num: int | None,
    c_desig: int | None,
    numeral: str,
    designation: str,
) -> None:
    """Write the row's numeral and designation, robust to both template layouts.

    When the numbered ("i)") and designation columns are distinct cells, the
    numeral goes in the narrow cell and the text in the wide one. When a single
    cell holds both (un-merged template variant), they are written together.
    The numeral is always the freshly computed one, so cloned rows can't keep a
    stale numeral copied from the template's last slot.
    """
    if c_desig is None or c_desig >= len(cells):
        return
    desig_cell = cells[c_desig]
    num_cell = cells[c_num] if (c_num is not None and c_num < len(cells)) else None
    if num_cell is not None and desig_cell is not num_cell:
        _set(num_cell, numeral)
        _set(desig_cell, designation)
        return
    _set(desig_cell, f"{numeral} {designation}".strip())


def _fill_basis(sec: SectionContext, table: Table) -> None:
    """Fill numbered category rows with designation, count and cover.

    The row layout differs per product, so we resolve target columns from the
    header labels rather than fixed indices.
    """
    headers = [_ctext(c).lower() for c in table.rows[0].cells]

    def col(*needles: str) -> int | None:
        for idx, h in enumerate(headers):
            if any(n in h for n in needles):
                return idx
        return None

    # "Category of Employees / Designation" is a header merged across two grid
    # columns: a narrow numbered column ("i)", "ii)") and the wide designation
    # column. Resolve both endpoints so the designation lands in the wide cell —
    # writing it into the narrow numbered cell wraps it one char per line.
    desig_cols = [
        i for i, h in enumerate(headers) if "category of employees" in h or "designation" in h
    ]
    c_num = desig_cols[0] if desig_cols else None
    c_desig = desig_cols[-1] if desig_cols else None
    c_count = col("no. of employees")
    c_plan = col("plan name")
    c_class = col("classification")
    c_rb = col("room & board", "room and board")
    c_basis = col("basis of cover", "sum insured")
    c_benefit = col("benefits provided")
    c_panel = col("panel")
    c_max = col("max limit")
    c_co = col("co-insurance")

    anchor_col = c_num if c_num is not None else 0
    # Group physical rows into category "units" (one numbered row + any
    # vertical-merge continuation rows below it), then clone the last unit until
    # every category has a row. This lets the form grow past its template's
    # i) to iv) slots instead of dropping the smaller cohorts.
    units = _basis_units(table, anchor_col)
    if not units:
        return
    needed = min(len(sec.basis_rows), MAX_BASIS_ROWS)
    while len(units) < needed:
        before = len(units)
        _clone_last_unit(table, units)
        units = _basis_units(table, anchor_col)
        if len(units) <= before:
            break  # safety: a clone that isn't re-detected as a unit would loop forever

    for idx, (unit, br) in enumerate(zip(units, sec.basis_rows, strict=False)):
        cells = table.rows[unit[0]].cells
        _write_designation(cells, c_num, c_desig, f"{_roman(idx + 1)})", br.designation)
        # Always write (even blank) so cloned rows don't inherit stale values.
        if c_count is not None and c_count < len(cells):
            _set(cells[c_count], br.num_employees)
        if c_plan is not None and c_plan < len(cells):
            _set(cells[c_plan], br.plan_name)
        if c_class is not None and c_class < len(cells):
            _set(cells[c_class], br.classification)
        if c_rb is not None and c_rb < len(cells):
            _set(cells[c_rb], br.room_board)
        if c_basis is not None and c_basis < len(cells):
            _set(cells[c_basis], br.sum_insured or br.room_board)
        if c_max is not None and c_max < len(cells):
            _set(cells[c_max], br.max_limit)
        if c_co is not None and c_co < len(cells):
            _set(cells[c_co], br.co_insurance)

        # The clinical page provides three physical benefit rows inside each
        # designation/plan unit. Populate them from that plan's configured
        # benefit schedule; never substitute a fixed GP/SP schedule.
        if c_benefit is not None:
            for line_idx, row_idx in enumerate(unit):
                line_cells = table.rows[row_idx].cells
                line = br.clinical_lines[line_idx] if line_idx < len(br.clinical_lines) else None
                for target, value in (
                    (c_benefit, line.name if line else ""),
                    (c_panel, line.panel if line else ""),
                    (c_max, line.max_limit if line else ""),
                    (c_co, line.co_insurance if line else ""),
                ):
                    if target is not None and target < len(line_cells):
                        _set(line_cells[target], value)

    total_row = next((row for row in table.rows if _ctext(row.cells[0]).lower() == "total"), None)
    if total_row is not None and c_count is not None:
        cells = _physical_cells(total_row)
        if len(cells) > 1:
            _set(cells[-1], sum(row.num_employees for row in sec.basis_rows[:needed]))


def _fill_age_band(sec: SectionContext, table: Table) -> None:
    if not sec.age_bands:
        return
    # Columns: find the Male / Female header cells (second header row usually).
    count_male = count_female = sum_male = sum_female = None
    if len(table.rows) > 1:
        labels = [_ctext(cell).lower() for cell in _physical_cells(table.rows[1])]
        male_cols = [idx for idx, label in enumerate(labels) if label == "male"]
        female_cols = [idx for idx, label in enumerate(labels) if label == "female"]
        if male_cols:
            count_male = male_cols[0]
            sum_male = male_cols[-1] if len(male_cols) > 1 else None
        if female_cols:
            count_female = female_cols[0]
            sum_female = female_cols[-1] if len(female_cols) > 1 else None
    if count_male is None or count_female is None:
        return
    band_labels = {lbl for lbl, _, _ in AGE_BANDS}
    for row in table.rows:
        label = _ctext(row.cells[0])
        if label in band_labels and label in sec.age_bands:
            cells = _physical_cells(row)
            m, f = sec.age_bands[label]
            _set(cells[count_male], m)
            _set(cells[count_female], f)
            sums = sec.age_band_sums.get(label)
            if sums and sum_male is not None and sum_female is not None:
                _set(cells[sum_male], _money(sums[0]))
                _set(cells[sum_female], _money(sums[1]))

    total_row = next((row for row in table.rows if _ctext(row.cells[0]).lower() == "total"), None)
    if total_row is not None:
        cells = _physical_cells(total_row)
        _set(cells[count_male], sum(value[0] for value in sec.age_bands.values()))
        _set(cells[count_female], sum(value[1] for value in sec.age_bands.values()))
        if sum_male is not None and sum_female is not None:
            _set(cells[sum_male], _money(sum(v[0] for v in sec.age_band_sums.values())))
            _set(cells[sum_female], _money(sum(v[1] for v in sec.age_band_sums.values())))


def _money(value: float) -> str:
    return f"{value:,.2f}" if value % 1 else f"{value:,.0f}"


def _compact_money(value: float) -> str:
    """Exact short notation for the form's unusually narrow amount cell."""
    for divisor, suffix in ((1_000_000, "m"), (1_000, "k")):
        scaled = value / divisor
        if value >= divisor and scaled.is_integer():
            return f"S${int(scaled)}{suffix}"
    return f"S${_money(value)}"


def _mark_amount_range(table: Table, amount: float) -> None:
    ranges = (
        (0, 200_000),
        (200_000, 500_000),
        (500_000, 1_000_000),
        (1_000_000, 2_000_000),
        (2_000_000, None),
    )
    label_rows = [1, 3, 5, 7, 9]
    for idx, (low, high) in enumerate(ranges):
        if amount >= low and (high is None or amount < high):
            cells = _physical_cells(table.rows[label_rows[idx]])
            if len(cells) >= 2:
                _mark(cells[-2])
            return


def _fill_highest_sum(sec: SectionContext, table: Table) -> None:
    if sec.highest_sum_insured is None:
        return
    _keep_table_together(table)
    if sec.highest_sum_insured_age is not None:
        cells = _physical_cells(table.rows[0])
        if len(cells) > 2:
            _set(cells[2], sec.highest_sum_insured_age)
    amount_cells = _physical_cells(table.rows[1])
    if len(amount_cells) > 3:
        _set_compact(amount_cells[2], f"Amt: {_compact_money(sec.highest_sum_insured)}")
        _set(amount_cells[3], "")
    _mark_amount_range(table, sec.highest_sum_insured)


def _fill_oldest_sum(sec: SectionContext, table: Table) -> None:
    if sec.oldest_insured_sum is None:
        return
    _keep_table_together(table)
    if sec.oldest_insured_age is not None:
        cells = _physical_cells(table.rows[0])
        if len(cells) > 2:
            _set(cells[2], sec.oldest_insured_age)
    amount_cells = _physical_cells(table.rows[1])
    if len(amount_cells) > 3:
        _set_compact(amount_cells[2], f"Amt: {_compact_money(sec.oldest_insured_sum)}")
        _set(amount_cells[3], "")
    _mark_amount_range(table, sec.oldest_insured_sum)


def _fill_family(sec: SectionContext, table: Table, data: dict[str, dict[str, int]]) -> None:
    """Family-composition matrix: Plan rows × (EO / ES / EC / EF)."""
    plan_names = list(dict.fromkeys([*sec.available_plans, *data]))
    if not plan_names:
        return
    # Map header columns to family buckets.
    bucket_col: dict[str, int] = {}
    for hrow in table.rows[:2]:
        for j, c in enumerate(hrow.cells):
            t = _ctext(c).lower()
            if "employee only" in t:
                bucket_col["EO"] = j
            elif "spouse" in t:
                bucket_col["ES"] = j
            elif "child" in t:
                bucket_col["EC"] = j
            elif "family" in t:
                bucket_col["EF"] = j
    if not bucket_col:
        return
    plan_rows = [r for r in table.rows if _ctext(r.cells[0]).lower().startswith("plan")]
    while len(plan_rows) < min(len(plan_names), MAX_BASIS_ROWS):
        table._tbl.append(copy.deepcopy(plan_rows[-1]._tr))
        plan_rows = [r for r in table.rows if _ctext(r.cells[0]).lower().startswith("plan")]
    for row, plan_name in zip(plan_rows, plan_names, strict=False):
        counts = data.get(plan_name, {})
        _set(row.cells[0], plan_name)
        for bucket, jcol in bucket_col.items():
            if jcol < len(row.cells):
                _set(row.cells[jcol], counts.get(bucket, 0))


def _combined_family(sec: SectionContext) -> dict[str, dict[str, int]]:
    combined: dict[str, dict[str, int]] = {}
    for source in (sec.family_local, sec.family_foreign):
        for plan, counts in source.items():
            target = combined.setdefault(plan, {"EO": 0, "ES": 0, "EC": 0, "EF": 0})
            for bucket in target:
                target[bucket] += counts.get(bucket, 0)
    return combined


def _fill_travel_basis(sec: SectionContext, table: Table, last_para: str) -> None:
    """Fill travel rows only when the platform identifies their frequency.

    The current product/category model has no trip frequency, trip duration,
    destination-area, or leisure-only fields. Putting all configured members
    into either the frequent or infrequent table would invent an answer, so the
    tables intentionally remain available for completion until those fields
    exist in platform data.
    """
    return


def _data_rows(table: Table, header_rows: int) -> list:
    return list(table.rows[header_rows:])


def _ensure_data_rows(table: Table, header_rows: int, needed: int) -> list:
    rows = _data_rows(table, header_rows)
    if not rows or len(_physical_cells(rows[-1])) <= 1:
        return rows
    while len(rows) < needed:
        table._tbl.append(copy.deepcopy(rows[-1]._tr))
        rows = _data_rows(table, header_rows)
    return rows


def _fill_claim_summaries(sec: SectionContext, table: Table) -> None:
    if not sec.claim_summaries or len(table.rows) < 3:
        return
    rows = _ensure_data_rows(table, 2, len(sec.claim_summaries))
    is_clinical = "specialists / diagnostic" in _header_text(table).lower()
    for row, summary in zip(rows, sec.claim_summaries, strict=False):
        cells = row.cells
        if len(_physical_cells(row)) <= 1 or len(cells) < 7:
            continue
        _set(cells[0], summary.period)
        _set(cells[1], summary.employees)
        _set(cells[2], summary.claimants)
        if is_clinical:
            _set(cells[3], _money(summary.paid_amount))
            _set(cells[4], _money(summary.outstanding_amount))
            _set(cells[5], _money(summary.secondary_paid_amount))
            _set(cells[6], _money(summary.secondary_outstanding_amount))
        else:
            _set(cells[3], summary.paid_count)
            _set(cells[4], _money(summary.paid_amount + summary.secondary_paid_amount))
            _set(cells[5], summary.outstanding_count)
            _set(
                cells[6],
                _money(summary.outstanding_amount + summary.secondary_outstanding_amount),
            )


def _fill_claim_details(sec: SectionContext, table: Table) -> None:
    if not sec.claim_details or len(table.rows) < 3:
        return
    rows = _ensure_data_rows(table, 2, len(sec.claim_details))
    for row, detail in zip(rows, sec.claim_details, strict=False):
        cells = row.cells
        if len(cells) == 4:
            _set(cells[0], detail.incurred_date)
            _set(cells[1], detail.nature)
            _set(cells[2], _money(detail.paid_amount) if detail.paid_amount else "")
            _set(
                cells[3],
                _money(detail.outstanding_amount) if detail.outstanding_amount else "",
            )
        elif len(cells) >= 8:  # GBT's expanded claim-experience layout
            paid = detail.paid_amount > 0
            outstanding = detail.outstanding_amount > 0
            _set(cells[0], detail.incurred_date)
            _set(cells[1], detail.nature)
            _set(cells[2], sec.employees_count)
            _set(cells[3], 1)
            _set(cells[4], 1 if paid else 0)
            _set(cells[5], _money(detail.paid_amount) if paid else "")
            _set(cells[6], 1 if outstanding else 0)
            _set(cells[7], _money(detail.outstanding_amount) if outstanding else "")


def generate(db: Session, policy_year: PolicyYear) -> tuple[bytes, list[str]]:
    """Build the context and render the filled ``.docx``; returns (bytes, notes)."""
    ctx = build_context(db, policy_year)
    return render_docx(ctx), ctx.completeness


__all__ = ["generate", "render_docx"]
