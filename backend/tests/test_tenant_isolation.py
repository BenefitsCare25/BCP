"""Cross-tenant isolation regression — every endpoint that takes a tenant-
owned resource ID must refuse access when the caller's `client_id` differs.

The test sets up TWO clients (A = seed demo, B = a second client created
inline) and confirms client A signed in cannot read or mutate client B's
policy year, categories, employees, dependants, audit log, or placement slips.

A user with no `client_id` (mock without a tenant binding) should also be
refused. System admins can read across tenants but every cross-tenant access
should be recorded in the audit log with `cross_tenant_access=True`.
"""
from __future__ import annotations

import os
from datetime import UTC, date, datetime
from pathlib import Path

import pytest

TEST_DB = Path(__file__).parent / "_test_tenant_isolation.db"
os.environ["INSPRO_DATABASE_URL"] = f"sqlite:///{TEST_DB}"

from fastapi.testclient import TestClient  # noqa: E402

from app.core.auth import (  # noqa: E402
    DEMO_BROKER_FIRM_ID,
    DEMO_CLIENT_ID,
    CurrentUser,
    get_current_user,
)
from app.db.base import Base  # noqa: E402
from app.db.session import SessionLocal, engine  # noqa: E402
from app.main import app  # noqa: E402
from app.models import (  # noqa: E402
    BulkPlanUpdate,
    Category,
    ClaimDocType,
    ClaimReviewConfig,
    Client,
    Dependant,
    Employee,
    EmployeeAttributeSchema,
    Enrollment,
    EnrollmentWindow,
    EntityAlias,
    Insurer,
    PanelCard,
    PanelListing,
    Plan,
    PolicyYear,
    Product,
)
from app.models.category import CategoryStatus, SourceKind  # noqa: E402
from app.models.policy_year import PolicyYearStatus  # noqa: E402
from scripts.seed_demo import seed  # noqa: E402

CLIENT_B_ID = "00000000-0000-0000-0000-0000000000b0"
USER_B_ID = "00000000-0000-0000-0000-0000000000b1"


def _user_a() -> CurrentUser:
    return CurrentUser(
        user_id="00000000-0000-0000-0000-000000000001",
        broker_firm_id=DEMO_BROKER_FIRM_ID,
        client_id=DEMO_CLIENT_ID,
        role="broker_admin",
    )


def _user_b() -> CurrentUser:
    return CurrentUser(
        user_id=USER_B_ID,
        broker_firm_id=DEMO_BROKER_FIRM_ID,
        client_id=CLIENT_B_ID,
        role="broker_admin",
    )


def _user_no_client() -> CurrentUser:
    return CurrentUser(
        user_id="00000000-0000-0000-0000-0000000000c0",
        broker_firm_id=DEMO_BROKER_FIRM_ID,
        client_id=None,
        role="broker_admin",
    )


def _user_system_admin() -> CurrentUser:
    return CurrentUser(
        user_id="00000000-0000-0000-0000-0000000000d0",
        broker_firm_id=DEMO_BROKER_FIRM_ID,
        client_id=DEMO_CLIENT_ID,
        role="system_admin",
    )


@pytest.fixture(scope="module", autouse=True)
def _setup_db():
    if TEST_DB.exists():
        TEST_DB.unlink()
    Base.metadata.create_all(bind=engine)
    seed()

    # Add a second client (B) with its own policy year + category + employee
    # so every endpoint has a B-owned resource to attempt to access from A.
    with SessionLocal() as session:
        ghs_product = session.query(Product).filter(Product.code == "GHS").first()
        assert ghs_product is not None
        py_a = (
            session.query(PolicyYear)
            .filter(PolicyYear.client_id == DEMO_CLIENT_ID)
            .first()
        )
        assert py_a is not None
        py_a.status = PolicyYearStatus.active
        session.add(
            Plan(
                id=PLAN_A_REVIEW,
                product_id=ghs_product.id,
                policy_year_id=py_a.id,
                code="A-GHS",
                display_name="Client A GHS plan",
                status="confirmed",
            )
        )

        client_b = Client(
            id=CLIENT_B_ID,
            name="Client B (test)",
            broker_firm_id=DEMO_BROKER_FIRM_ID,
        )
        session.add(client_b)
        session.flush()

        py_b = PolicyYear(
            id="00000000-0000-0000-0000-0000000000b2",
            client_id=CLIENT_B_ID,
            year=2026,
            start_date=date(2026, 1, 1),
            end_date=date(2026, 12, 31),
            status=PolicyYearStatus.active,
        )
        session.add(py_b)
        session.flush()

        cat_b = Category(
            id="00000000-0000-0000-0000-0000000000b3",
            policy_year_id=py_b.id,
            priority=1,
            display_name="Client B secret category",
            raw_description="Client B secret category",
            source=SourceKind.system_generated.value,
            status=CategoryStatus.needs_review.value,
            human_modified=False,
        )
        session.add(cat_b)

        emp_b = Employee(
            id="00000000-0000-0000-0000-0000000000b4",
            client_id=CLIENT_B_ID,
            policy_year_id=py_b.id,
            staff_id="B-STAFF-1",
            employee_name="Bea Beta",
            attribute_values={"grade": 18},
            derived_attribute_values={},
            source="csv_import",
            status="active",
        )
        session.add(emp_b)
        session.flush()

        dep_b = Dependant(
            id="00000000-0000-0000-0000-0000000000b5",
            client_id=CLIENT_B_ID,
            policy_year_id=py_b.id,
            employee_id=emp_b.id,
            attribute_values={"relationship": "spouse"},
            link_method="staff_id",
            status="active",
        )
        session.add(dep_b)

        # A plan owned by client B (for cross-tenant plan access checks).
        product_id = ghs_product.id
        plan_b = Plan(
            id=PLAN_B,
            product_id=product_id,
            policy_year_id=py_b.id,
            code="B1",
            display_name="Client B plan",
            status="needs_review",
        )
        session.add(plan_b)

        # An enrollment window owned by client B (for window-id isolation checks).
        win_b = EnrollmentWindow(
            id=WINDOW_B,
            policy_year_id=py_b.id,
            client_id=CLIENT_B_ID,
            name="Client B window",
            opens_at=datetime(2026, 1, 1, tzinfo=UTC),
            closes_at=datetime(2026, 2, 1, tzinfo=UTC),
        )
        session.add(win_b)
        session.flush()

        # An enrollment owned by client B (for enrollment-id isolation checks).
        enr_b = Enrollment(
            id=ENROLL_B,
            window_id=win_b.id,
            policy_year_id=py_b.id,
            client_id=CLIENT_B_ID,
            employee_id=emp_b.id,
        )
        session.add(enr_b)

        # A bulk coverage batch owned by client B (history/detail/undo checks).
        session.add(
            BulkPlanUpdate(
                id=BULK_B,
                policy_year_id=py_b.id,
                client_id=CLIENT_B_ID,
                product_code="MED",
                action="set_plan",
                target_plan_code="B1",
                selector={"employee_ids": [EMP_B]},
                result_summary={"counts": {"applied": 0}, "restore": []},
            )
        )

        # A panel listing owned by client B (clinic-locator isolation checks).
        session.add(
            PanelListing(
                id=PANEL_B,
                client_id=CLIENT_B_ID,
                insurer="GE-SG",
                panel_provider="Adept",
                country="SG",
                clinic_type="gp",
            )
        )

        # A panel e-card PINNED to client B (e-card isolation checks).
        session.add(
            PanelCard(
                id=CARD_B,
                client_id=CLIENT_B_ID,
                insurer="GE-SG",
                panel_provider="Adept",
                name="GE Adept Card",
                artwork_front_path="firm/b/panel_card/x/front.png",
            )
        )

        # An insurer catalog entry owned by client B (name-catalog isolation).
        session.add(
            Insurer(
                id=INSURER_B,
                client_id=CLIENT_B_ID,
                name="Client B Only Insurer",
                legal_name="Client B Only Insurer Pte Ltd",
            )
        )

        # An entity alias owned by client B (alias-map isolation).
        session.add(
            EntityAlias(
                id=ALIAS_B,
                client_id=CLIENT_B_ID,
                alias="BSUB",
                canonical="Client B Subsidiary Pte Ltd",
                alias_normalized="bsub",
            )
        )

        # A claim document-type row owned by client B (registry isolation).
        session.add(
            ClaimDocType(
                id=DOCTYPE_B,
                client_id=CLIENT_B_ID,
                key="client_b_doc",
                display="Client B Document",
                aliases=["client b document"],
                key_fields=[{"name": "Secret Field", "keywords": []}],
            )
        )

        # A per-claim-type review rule setup owned by client B.
        session.add(
            ClaimReviewConfig(
                id=REVIEW_CONFIG_B,
                client_id=CLIENT_B_ID,
                claim_kind="insured",
                claim_key="GHS",
                display_label="Client B GHS rules",
                field_maps=[
                    {
                        "portal_field": "amount_claimed",
                        "document_field": "Total Amount",
                        "mode": "numeric",
                        "tolerance": 0.01,
                        "verify_with_vision": True,
                    }
                ],
                ai_rules=[
                    {
                        "id": "rule_1",
                        "rule": "Client B secret rule",
                        "category": "general",
                        "severity": "critical",
                    }
                ],
            )
        )
        for config_id, scope_code, label in (
            (
                REVIEW_CONFIG_B_GOVT,
                "ghs_hospitalisation_govt",
                "Client B hospitalisation rules",
            ),
            (
                REVIEW_CONFIG_B_PRIVATE,
                "ghs_hospitalisation_private",
                "Client B hospitalisation rules",
            ),
            (
                REVIEW_CONFIG_B_UNAVAILABLE,
                "*",
                "Client B unavailable product rules",
            ),
        ):
            session.add(
                ClaimReviewConfig(
                    id=config_id,
                    client_id=CLIENT_B_ID,
                    claim_kind="insured",
                    claim_key=(
                        "BPROD"
                        if config_id == REVIEW_CONFIG_B_UNAVAILABLE
                        else "GHS"
                    ),
                    scope_code=scope_code,
                    display_label=label,
                    field_maps=[
                        {
                            "portal_field": "amount_claimed",
                            "document_field": "Total Amount",
                            "mode": "numeric",
                            "tolerance": 0.01,
                            "verify_with_vision": True,
                        }
                    ],
                    ai_rules=[],
                )
            )

        # A product + employee-attribute owned by client B (schemas CRUD
        # isolation — load_editable_global 404s on another tenant's row).
        session.add(
            Product(
                id=PRODUCT_B_OWNED,
                client_id=CLIENT_B_ID,
                code="BPROD",
                display_name="Client B product",
            )
        )
        session.add(
            EmployeeAttributeSchema(
                id=ATTR_B,
                client_id=CLIENT_B_ID,
                attribute_id="client_b_attr",
                display_name="Client B attribute",
                data_type="text",
            )
        )

        session.commit()

    yield
    engine.dispose()
    if TEST_DB.exists():
        TEST_DB.unlink()


@pytest.fixture
def client_as_a() -> TestClient:
    app.dependency_overrides[get_current_user] = _user_a
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.pop(get_current_user, None)


@pytest.fixture
def client_as_no_tenant() -> TestClient:
    app.dependency_overrides[get_current_user] = _user_no_client
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.pop(get_current_user, None)


@pytest.fixture
def client_as_system_admin() -> TestClient:
    app.dependency_overrides[get_current_user] = _user_system_admin
    try:
        yield TestClient(app)
    finally:
        app.dependency_overrides.pop(get_current_user, None)


# ── Client B resource IDs (created in fixture) ──────────────────────────────
PY_B = "00000000-0000-0000-0000-0000000000b2"
CAT_B = "00000000-0000-0000-0000-0000000000b3"
EMP_B = "00000000-0000-0000-0000-0000000000b4"
DEP_B = "00000000-0000-0000-0000-0000000000b5"
PLAN_B = "00000000-0000-0000-0000-0000000000b6"
PLAN_A_REVIEW = "00000000-0000-0000-0000-0000000000a6"
# Any product id — the cross-tenant guard rejects at the policy year before the
# product is ever looked up, so this need not exist.
PRODUCT_B = "00000000-0000-0000-0000-0000000000b7"
WINDOW_B = "00000000-0000-0000-0000-0000000000b8"
ENROLL_B = "00000000-0000-0000-0000-0000000000b9"
PANEL_B = "00000000-0000-0000-0000-0000000000ba"
INSURER_B = "00000000-0000-0000-0000-0000000000bc"
ALIAS_B = "00000000-0000-0000-0000-0000000000bd"
DOCTYPE_B = "00000000-0000-0000-0000-0000000000be"
CARD_B = "00000000-0000-0000-0000-0000000000bb"
# Any assignment id — the guard rejects at the policy year first.
CARD_ASSIGNMENT_B = "00000000-0000-0000-0000-0000000000bc"
PRODUCT_B_OWNED = "00000000-0000-0000-0000-0000000000bf"
ATTR_B = "00000000-0000-0000-0000-0000000000c1"
REVIEW_CONFIG_B = "00000000-0000-0000-0000-0000000000c2"
REVIEW_CONFIG_B_GOVT = "00000000-0000-0000-0000-0000000000c4"
REVIEW_CONFIG_B_PRIVATE = "00000000-0000-0000-0000-0000000000c5"
REVIEW_CONFIG_B_UNAVAILABLE = "00000000-0000-0000-0000-0000000000c6"
BULK_B = "00000000-0000-0000-0000-0000000000c3"


# ── PolicyYear ──────────────────────────────────────────────────────────────
def test_policy_year_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}")
    assert res.status_code == 404


def test_policy_year_set_current_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/policy-years/{PY_B}/set-current")
    assert res.status_code == 404


def test_policy_year_update_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/policy-years/{PY_B}", json={"claim_grace_period_days": 30}
    )
    assert res.status_code == 404


def test_policy_year_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/policy-years/{PY_B}")
    assert res.status_code == 404


def test_policy_year_copy_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/copy",
        json={"start_date": "2029-01-01", "end_date": "2029-12-31"},
    )
    assert res.status_code == 404


def test_policy_year_snapshot_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/snapshot")
    assert res.status_code == 404


def test_reports_benefit_selection_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/benefit-selection")
    assert res.status_code == 404


def test_reports_employee_listing_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/employee-listing?insurer=AIA"
    )
    assert res.status_code == 404


def test_reports_dependant_listing_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/dependant-listing?insurer=AIA"
    )
    assert res.status_code == 404


def test_reports_placement_slip_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/placement-slip")
    assert res.status_code == 404


def test_reports_quotation_slip_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/quotation-slip")
    assert res.status_code == 404


def test_claims_register_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/claims/register?policy_year_id={PY_B}")
    assert res.status_code == 404


def test_reports_readiness_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/readiness")
    assert res.status_code == 404


def test_reports_built_in_employee_listing_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/workbooks/member-register"
    )
    assert res.status_code == 404


def test_reports_built_in_dependant_listing_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/workbooks/member-register"
    )
    assert res.status_code == 404


def test_reports_portal_activity_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/activity-access")
    assert res.status_code == 404


def test_reports_company_activity_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/activity-access")
    assert res.status_code == 404


def test_reports_portal_access_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/activity-access")
    assert res.status_code == 404


# ── Report versions ──────────────────────────────────────────────────────────
REPORT_VERSION_B = "00000000-0000-0000-0000-0000000000bd"


def _ensure_report_version_b() -> None:
    """Idempotently create a report version owned by client B."""
    from app.models import ReportVersion

    with SessionLocal() as s:
        if s.get(ReportVersion, REPORT_VERSION_B) is None:
            s.add(
                ReportVersion(
                    id=REPORT_VERSION_B,
                    client_id=CLIENT_B_ID,
                    policy_year_id=PY_B,
                    report_type="employee_listing",
                    scope_key="testsure",
                    version_no=1,
                    mode="versioned",
                    params={"insurer": "TestSure", "masked": True},
                    summary={},
                    file_name="employee-listing-v1.xlsx",
                    mime_type="application/octet-stream",
                    size_bytes=1,
                    sha256="0" * 64,
                    storage_path="nofirm/b/report_version/x/y.xlsx",
                )
            )
            s.commit()


def test_report_versions_create_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/report-versions",
        json={"report_type": "placement_slip"},
    )
    assert res.status_code == 404


def test_report_versions_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/report-versions",
        params={"report_type": "placement_slip"},
    )
    assert res.status_code == 404


def test_report_versions_status_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/report-versions/status",
        params={"report_type": "placement_slip"},
    )
    assert res.status_code == 404


def test_report_version_download_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_report_version_b()
    res = client_as_a.get(f"/api/v1/report-versions/{REPORT_VERSION_B}/download")
    assert res.status_code == 404


def test_report_version_movement_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_report_version_b()
    res = client_as_a.get(f"/api/v1/report-versions/{REPORT_VERSION_B}/movement")
    assert res.status_code == 404


def test_report_version_movement_summary_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    """The counts endpoint is ungated for masking (three integers leak no
    identifier) but must still be tenant-scoped like every other resource."""
    _ensure_report_version_b()
    res = client_as_a.get(
        f"/api/v1/report-versions/{REPORT_VERSION_B}/movement-summary"
    )
    assert res.status_code == 404


def test_reports_member_listing_template_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/member-listing-template"
    )
    assert res.status_code == 404


def test_reports_underwriting_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/underwriting")
    assert res.status_code == 404


def test_underwriting_cases_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/underwriting/cases")
    assert res.status_code == 404


def test_underwriting_refresh_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/policy-years/{PY_B}/underwriting/refresh")
    assert res.status_code == 404


def test_underwriting_review_cross_tenant_404(client_as_a: TestClient) -> None:
    # A review belonging to tenant B must 404 for tenant A (user_owns →
    # _deny_cross_tenant), indistinguishable from a non-existent id.
    from app.models import UnderwritingReview

    with SessionLocal() as s:
        review = UnderwritingReview(
            client_id=CLIENT_B_ID, policy_year_id=PY_B, insurer="TestSure",
            employee_id=None, dependant_id=None,
        )
        s.add(review)
        s.commit()
        review_id = review.id
    try:
        res = client_as_a.patch(
            f"/api/v1/underwriting/reviews/{review_id}",
            json={"status": "completed"},
        )
        assert res.status_code == 404
    finally:
        with SessionLocal() as s:
            row = s.get(UnderwritingReview, review_id)
            if row is not None:
                s.delete(row)
                s.commit()


def test_fact_find_form_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/fact-find-form")
    assert res.status_code == 404


def test_fact_find_form_download_own_tenant(client_as_a: TestClient) -> None:
    import io

    from docx import Document

    rows = client_as_a.get("/api/v1/policy-years").json()
    assert rows, "expected at least one policy year for client A"
    res = client_as_a.get(f"/api/v1/policy-years/{rows[0]['id']}/fact-find-form")
    assert res.status_code == 200
    assert res.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert res.content[:2] == b"PK"  # valid .docx (zip) magic
    assert "x-factfind-notes" in {k.lower() for k in res.headers}
    # The filled template opens and the company-name cell is populated.
    doc = Document(io.BytesIO(res.content))
    cells = [
        p.text
        for t in doc.tables
        for r in t.rows
        for c in r.cells
        for p in c.paragraphs
    ]
    assert any("Name of Company" in c for c in cells)


def test_policy_year_list_only_returns_own_tenant(client_as_a: TestClient) -> None:
    rows = client_as_a.get("/api/v1/policy-years").json()
    ids = {r["id"] for r in rows}
    assert PY_B not in ids


# ── Product terms (per-product coverage periods) ────────────────────────────
def test_product_terms_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/product-terms")
    assert res.status_code == 404


def test_product_term_set_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/product-terms/{PRODUCT_B}",
        json={"coverage_start": "2026-04-01", "coverage_end": "2027-03-31"},
    )
    assert res.status_code == 404


def test_product_term_reset_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(
        f"/api/v1/policy-years/{PY_B}/product-terms/{PRODUCT_B}"
    )
    assert res.status_code == 404


def test_remove_product_cross_tenant_404(client_as_a: TestClient) -> None:
    # The policy year belongs to tenant B — load_policy_year must 404 before any
    # delete touches B's products.
    res = client_as_a.delete(f"/api/v1/policy-years/{PY_B}/products/GHS")
    assert res.status_code == 404


# ── Flex scheme (flexible benefits) ─────────────────────────────────────────
def test_flex_scheme_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/flex-scheme")
    assert res.status_code == 404


def test_flex_scheme_put_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/flex-scheme",
        json={"scheme": {"meta": {"currency": "SGD"}, "tiers": []}},
    )
    assert res.status_code == 404


def test_flex_scheme_confirm_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/policy-years/{PY_B}/flex-scheme/confirm")
    assert res.status_code == 404


def test_flex_scheme_extract_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/flex-scheme/extract",
        files={"files": ("doc.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )
    assert res.status_code == 404


def test_flex_scheme_membership_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/flex-scheme/membership")
    assert res.status_code == 404


def test_flex_scheme_assign_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/policy-years/{PY_B}/flex-scheme/assign")
    assert res.status_code == 404


def test_flex_scheme_coverage_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/flex-scheme/coverage")
    assert res.status_code == 404


def test_flex_scheme_coverage_export_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/flex-scheme/coverage/export")
    assert res.status_code == 404


def test_flex_scheme_roster_vocab_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/flex-scheme/roster-vocab")
    assert res.status_code == 404


def test_flex_scheme_suggest_matches_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/flex-scheme/suggest-matches"
    )
    assert res.status_code == 404


def test_flex_pricing_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/flex-pricing")
    assert res.status_code == 404


def test_flex_pricing_put_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/flex-pricing",
        json={"pricing": {"products": {}}},
    )
    assert res.status_code == 404


def test_voluntary_rates_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/products/{PRODUCT_B}/voluntary-rates"
    )
    assert res.status_code == 404


def test_voluntary_rates_put_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/products/{PRODUCT_B}/voluntary-rates",
        json={"bands": [{"label": "all", "min": None, "max": None, "rate": 1.0}]},
    )
    assert res.status_code == 404


# ── Categories ──────────────────────────────────────────────────────────────
def test_category_list_cross_tenant_policy_year_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/categories", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_category_grouped_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/categories/grouped", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_category_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/categories/{CAT_B}")
    assert res.status_code == 404


def test_category_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/categories/{CAT_B}", json={"display_name": "hijack"}
    )
    assert res.status_code == 404


def test_category_create_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        "/api/v1/categories",
        json={"policy_year_id": PY_B, "display_name": "intruder"},
    )
    assert res.status_code == 404


def test_category_confirm_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/categories/{CAT_B}/confirm")
    assert res.status_code == 404


def test_category_bulk_confirm_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post("/api/v1/categories/bulk-confirm", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_category_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/categories/{CAT_B}")
    assert res.status_code == 404


def test_category_bulk_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete("/api/v1/categories", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_category_ai_suggest_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/categories/{CAT_B}/ai-suggest")
    assert res.status_code == 404


def test_category_coverage_stats_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/categories/stats/coverage", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_placement_slips_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/placement-slips", params={"policy_year_id": PY_B})
    assert res.status_code == 404


# ── Employees / Dependants / Matches ────────────────────────────────────────
def test_employee_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/employees", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_employee_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}")
    assert res.status_code == 404


def test_employee_benefit_statement_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/benefit-statement")
    assert res.status_code == 404


def test_employee_coverage_summary_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        "/api/v1/employees/coverage-summary", params={"policy_year_id": PY_B}
    )
    assert res.status_code == 404


def test_employee_coverage_summary_leavers_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    """`include_left` widens WHICH rows come back, never whose. A filter that
    relaxes a WHERE clause is exactly where a tenant check gets forgotten."""
    res = client_as_a.get(
        "/api/v1/employees/coverage-summary",
        params={"policy_year_id": PY_B, "include_left": "true"},
    )
    assert res.status_code == 404


def test_employee_coverage_report_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        "/api/v1/employees/coverage-report/export", params={"policy_year_id": PY_B}
    )
    assert res.status_code == 404


def test_dependant_coverage_report_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        "/api/v1/dependants/coverage-report/export", params={"policy_year_id": PY_B}
    )
    assert res.status_code == 404


_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_FAKE_XLSX = ("adc.xlsx", b"PK\x03\x04 not-a-real-xlsx", _XLSX_MIME)


def test_member_listing_template_cross_tenant_404(client_as_a: TestClient) -> None:
    # This replaced /adc/template as the roster download, and it carries the
    # WHOLE listing with unmasked identifiers — so it is the surface that has to
    # be tenant-guarded. Pointing this test at the deleted route would have
    # passed on routing alone, proving nothing.
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/member-listing-template"
    )
    assert res.status_code == 404


def test_adc_preview_cross_tenant_404(client_as_a: TestClient) -> None:
    # The tenant guard fires before the upload is parsed, so a fake file is fine.
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/adc/preview", files={"file": _FAKE_XLSX}
    )
    assert res.status_code == 404


def test_adc_apply_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/adc/apply", files={"file": _FAKE_XLSX}
    )
    assert res.status_code == 404


def test_employee_portal_preview_cross_tenant_404(client_as_a: TestClient) -> None:
    for suffix in (
        "",
        "/benefit-statement",
        "/utilization",
        "/coverage-options",
        "/enrollment",
        "/dependants",
        "/claims",
    ):
        res = client_as_a.get(f"/api/v1/employees/{EMP_B}/portal-preview{suffix}")
        assert res.status_code == 404, suffix


def test_employee_bulk_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete("/api/v1/employees", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_employee_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/employees/{EMP_B}", json={"employee_name": "hijack"}
    )
    assert res.status_code == 404


# ── Plan overrides (enrollment module) ──────────────────────────────────────
def test_plan_overrides_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/plan-overrides")
    assert res.status_code == 404


def test_plan_overrides_put_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/employees/{EMP_B}/plan-overrides/GHS",
        json={"plan_code": "B1"},
    )
    assert res.status_code == 404


def test_plan_overrides_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/employees/{EMP_B}/plan-overrides/GHS")
    assert res.status_code == 404


def test_coverage_history_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/coverage-history")
    assert res.status_code == 404


def test_coverage_revert_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/employees/{EMP_B}/coverage/revert", json={"target": "default"}
    )
    assert res.status_code == 404


def test_enrollment_reset_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/enrollments/{ENROLL_B}/reset")
    assert res.status_code == 404


def test_enrollment_reopen_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/enrollments/{ENROLL_B}/reopen")
    assert res.status_code == 404


# ── Enrollment windows + leave policy (enrollment module) ───────────────────
def test_enrollment_windows_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/enrollment-windows")
    assert res.status_code == 404


def test_enrollment_window_create_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/enrollment-windows",
        json={"name": "x", "opens_at": "2026-01-01T00:00:00Z",
              "closes_at": "2026-02-01T00:00:00Z"},
    )
    assert res.status_code == 404


def test_enrollment_window_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/enrollment-windows/{WINDOW_B}")
    assert res.status_code == 404


def test_enrollment_window_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/enrollment-windows/{WINDOW_B}", json={"name": "hijack"}
    )
    assert res.status_code == 404


def test_enrollment_window_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/enrollment-windows/{WINDOW_B}")
    assert res.status_code == 404


def test_leave_policy_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/leave-policy")
    assert res.status_code == 404


def test_leave_policy_put_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/leave-policy", json={"allow_buy": True}
    )
    assert res.status_code == 404


def test_leave_rate_options_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/leave-rate-options")
    assert res.status_code == 404


def test_window_open_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/enrollment-windows/{WINDOW_B}/open")
    assert res.status_code == 404


def test_window_close_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/enrollment-windows/{WINDOW_B}/close")
    assert res.status_code == 404


def test_window_enrollments_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/enrollment-windows/{WINDOW_B}/enrollments")
    assert res.status_code == 404


def test_enrollment_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/enrollments/{ENROLL_B}")
    assert res.status_code == 404


def test_enrollment_options_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/enrollments/{ENROLL_B}/options")
    assert res.status_code == 404


def test_enrollment_elections_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/enrollments/{ENROLL_B}/elections",
        json={"elections": [{"product_code": "MED", "plan_code": "B1"}]},
    )
    assert res.status_code == 404


def test_enrollment_leave_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/enrollments/{ENROLL_B}/leave", json={"action": "none", "days": 0}
    )
    assert res.status_code == 404


def test_enrollment_submit_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/enrollments/{ENROLL_B}/submit")
    assert res.status_code == 404


def test_enrollment_confirm_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/enrollments/{ENROLL_B}/confirm")
    assert res.status_code == 404


def test_bulk_preview_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/bulk-plan-updates/preview",
        json={"product_code": "MED", "action": "decline",
              "selector": {"employee_ids": [EMP_B]}},
    )
    assert res.status_code == 404


def test_bulk_apply_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/bulk-plan-updates/apply",
        json={"product_code": "MED", "action": "decline",
              "selector": {"employee_ids": [EMP_B]}},
    )
    assert res.status_code == 404


def test_bulk_history_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/bulk-plan-updates")
    assert res.status_code == 404


def test_bulk_detail_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/bulk-plan-updates/{BULK_B}")
    assert res.status_code == 404


def test_bulk_undo_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/bulk-plan-updates/{BULK_B}/undo")
    assert res.status_code == 404


def test_member_facets_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/member-facets")
    assert res.status_code == 404


def test_member_query_count_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/member-query/count",
        json={"query": {"employee_ids": [EMP_B]}},
    )
    assert res.status_code == 404


def test_member_query_list_cross_tenant_404(client_as_a: TestClient) -> None:
    # An EMPTY body is the listing's default view, so this endpoint would happily
    # return another tenant's whole roster if the year guard ever came off.
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/member-query/list",
        json={},
    )
    assert res.status_code == 404


def test_member_query_resolve_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/member-query/resolve",
        json={"text": "B-1"},
    )
    assert res.status_code == 404


def test_dependant_facets_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/dependant-facets")
    assert res.status_code == 404


def test_dependant_query_list_cross_tenant_404(client_as_a: TestClient) -> None:
    # As with the member listing, an EMPTY body is the default view — so a
    # missing year guard would hand over another tenant's dependants wholesale.
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/dependant-query/list",
        json={},
    )
    assert res.status_code == 404


def test_dual_coverage_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/dual-coverage")
    assert res.status_code == 404


def test_dual_coverage_decision_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/dual-coverage/decisions",
        json={"subject_key": "whatever", "decision": "intentional_both"},
    )
    assert res.status_code == 404
    res = client_as_a.delete(
        f"/api/v1/policy-years/{PY_B}/dual-coverage/decisions/whatever"
    )
    assert res.status_code == 404


def test_dual_coverage_set_cover_cross_tenant_404(client_as_a: TestClient) -> None:
    """The one endpoint here that MOVES MONEY — it writes plan overrides, so a
    cross-tenant hit would change another firm's premiums, not just read them."""
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/dual-coverage/dependants/whatever/cover",
        json={"covered": False},
    )
    assert res.status_code == 404


def test_dual_coverage_decisions_is_a_tenant_table() -> None:
    """A model missing from `models/__init__` is invisible to Alembic
    autogenerate AND to `sync_firm_schema` — so on Postgres every firm's queries
    fall through `search_path` to the `public` copy and SILENTLY SHARE ROWS.
    No error, just cross-tenant data. This is the cheap guard against that."""
    import app.models  # noqa: F401 — registers every model in Base.metadata
    from app.db.tenancy import CONTROL_TABLES, tenant_tables

    assert "dual_coverage_decisions" in {t.name for t in tenant_tables()}
    assert "dual_coverage_decisions" not in CONTROL_TABLES


def test_orphan_overrides_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/plan-overrides/orphans")
    assert res.status_code == 404


def test_dependant_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/dependants/{DEP_B}", json={"attribute_values": {"x": 1}}
    )
    assert res.status_code == 404


def test_dependant_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/dependants", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_dependant_bulk_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete("/api/v1/dependants", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_match_results_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/match-results", params={"policy_year_id": PY_B})
    assert res.status_code == 404


def test_match_run_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post("/api/v1/match-results/run", params={"policy_year_id": PY_B})
    assert res.status_code == 404


# ── Plans (Schedule of Benefits) ─────────────────────────────────────────────
def test_plan_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/plans/{PLAN_B}")
    assert res.status_code == 404


def test_plan_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/plans/{PLAN_B}", json={"display_name": "hijack"}
    )
    assert res.status_code == 404


# ── Config recommendations ──────────────────────────────────────────────────
def test_recommend_config_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(f"/api/v1/policy-years/{PY_B}/recommend-config")
    assert res.status_code == 404


def test_apply_config_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/apply-config",
        json={"attributes": [{"attribute_id": "hijack", "display_name": "x",
                              "data_type": "string"}], "products": [],
              "rerun_matching": False},
    )
    assert res.status_code == 404


# ── Product setup ───────────────────────────────────────────────────────────
def test_product_setups_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/product-setups")
    assert res.status_code == 404


def test_product_setup_get_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/product-setups/GHS")
    assert res.status_code == 404


def test_product_setup_save_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/product-setups/GHS",
        json={"answers": {}, "template_version": 1},
    )
    assert res.status_code == 404


def test_product_setup_confirm_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/product-setups/GHS/confirm",
        json={"answers": {"plans": [{"code": "1", "selected": True}]},
              "template_version": 1},
    )
    assert res.status_code == 404


def test_product_setup_discard_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/policy-years/{PY_B}/product-setups/GHS")
    assert res.status_code == 404


def test_product_setup_stale_conflict_is_structured(
    client_as_a: TestClient,
) -> None:
    policy_year_id = client_as_a.get("/api/v1/policy-years").json()[0]["id"]
    path = f"/api/v1/policy-years/{policy_year_id}/product-setups/GTPD"
    created = client_as_a.put(
        path,
        json={"answers": {}, "template_version": 1},
    )
    assert created.status_code == 200, created.text

    stale = client_as_a.put(
        path,
        json={"answers": {"plans": []}, "template_version": 1},
    )
    assert stale.status_code == 409
    assert stale.json()["detail"]["code"] == "stale_configuration"
    assert "Reload the latest version" in stale.json()["detail"]["message"]

    assert client_as_a.delete(path).status_code == 204


def test_member_counts_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/member-counts",
        json={"product_code": "GHS", "has_dependants": True, "categories": []},
    )
    assert res.status_code == 404


# ── Audit log ───────────────────────────────────────────────────────────────
def test_audit_log_scoped_to_own_tenant(client_as_a: TestClient) -> None:
    """Audit log returns only the caller's own client rows."""
    res = client_as_a.get("/api/v1/audit-log")
    assert res.status_code == 200
    rows = res.json()["items"]
    foreign = [r for r in rows if r.get("client_id") and r["client_id"] != DEMO_CLIENT_ID]
    assert not foreign, f"audit log leaked rows from other tenants: {foreign}"


def test_audit_log_system_admin_sees_all(client_as_system_admin: TestClient) -> None:
    """System admins are exempt — they see every tenant's rows."""
    res = client_as_system_admin.get("/api/v1/audit-log")
    assert res.status_code == 200


# ── User without a client_id ────────────────────────────────────────────────
def test_no_tenant_user_cannot_list_policy_years(client_as_no_tenant: TestClient) -> None:
    res = client_as_no_tenant.get("/api/v1/policy-years")
    assert res.status_code == 400
    assert "active client" in res.text.lower()


# ── Slip template profiles (broker-corrected column mappings) ────────────────
def _put_profile_as(user_factory, payload: dict):
    """PUT the template-profile endpoint as a specific user. Sets the override
    per-call (the shared client_as_* fixtures can't both be active at once)."""
    app.dependency_overrides[get_current_user] = user_factory
    try:
        return TestClient(app).put(
            "/api/v1/placement-slips/template-profiles", json=payload
        )
    finally:
        app.dependency_overrides.pop(get_current_user, None)


def test_template_profile_is_scoped_per_tenant() -> None:
    """A and B saving the SAME fingerprint must not collide — the override is
    keyed by the caller's active client, so neither can read or overwrite the
    other's mapping."""
    fp = "shared-fp-xyz"
    a = _put_profile_as(
        _user_a, {"fingerprint": fp, "product_code": "GBT", "roles": {"name_col": 0}}
    )
    assert a.status_code == 200, a.text
    b = _put_profile_as(
        _user_b, {"fingerprint": fp, "product_code": "GBT", "roles": {"name_col": 3}}
    )
    assert b.status_code == 200, b.text

    # Two distinct rows (one per tenant), each with its own mapping.
    from app.models import SlipTemplateProfile

    with SessionLocal() as s:
        a_row = (
            s.query(SlipTemplateProfile)
            .filter_by(client_id=DEMO_CLIENT_ID, fingerprint=fp)
            .one()
        )
        b_row = (
            s.query(SlipTemplateProfile)
            .filter_by(client_id=CLIENT_B_ID, fingerprint=fp)
            .one()
        )
    assert a_row.id != b_row.id
    assert a_row.roles["name_col"] == 0  # A's mapping untouched by B's save
    assert b_row.roles["name_col"] == 3


def test_template_profile_save_requires_tenant() -> None:
    res = _put_profile_as(
        _user_no_client,
        {"fingerprint": "fp", "product_code": "GBT", "roles": {"name_col": 0}},
    )
    assert res.status_code == 400
    assert "active client" in res.text.lower()


# ── Member accounts (employee-portal provisioning) ──────────────────────────
MEMBER_ACC_B = "00000000-0000-0000-0000-0000000000ba"


def _ensure_member_account_b() -> None:
    """Idempotently create a portal member account owned by client B."""
    from app.models import MemberAccount

    with SessionLocal() as s:
        if s.get(MemberAccount, MEMBER_ACC_B) is None:
            s.add(
                MemberAccount(
                    id=MEMBER_ACC_B,
                    client_id=CLIENT_B_ID,
                    email="member@client-b.test",
                    staff_id="B-STAFF-1",
                    status="active",
                )
            )
            s.commit()


def test_member_account_create_for_b_employee_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/employees/{EMP_B}/member-account", json={"email": "x@y.test"}
    )
    assert res.status_code == 404


def test_member_account_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    _ensure_member_account_b()
    res = client_as_a.get("/api/v1/member-accounts")
    assert res.status_code == 200
    ids = {item["id"] for item in res.json()["items"]}
    assert MEMBER_ACC_B not in ids


def test_member_account_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_member_account_b()
    res = client_as_a.patch(
        f"/api/v1/member-accounts/{MEMBER_ACC_B}", json={"status": "disabled"}
    )
    assert res.status_code == 404


def test_member_account_resend_invite_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_member_account_b()
    res = client_as_a.post(f"/api/v1/member-accounts/{MEMBER_ACC_B}/resend-invite")
    assert res.status_code == 404


def test_member_account_bulk_invite_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        "/api/v1/member-accounts/bulk-invite", json={"policy_year_id": PY_B}
    )
    assert res.status_code == 404


# ── Claims (employee-portal claims module) ───────────────────────────────────
CLAIM_B = "00000000-0000-0000-0000-0000000000bc"
ENQUIRY_B = "00000000-0000-0000-0000-0000000000q9"


def _ensure_claim_b() -> None:
    """Idempotently create a claim owned by client B."""
    from datetime import date as _date

    from app.models import Claim

    with SessionLocal() as s:
        if s.get(Claim, CLAIM_B) is None:
            s.add(
                Claim(
                    id=CLAIM_B,
                    client_id=CLIENT_B_ID,
                    policy_year_id=PY_B,
                    employee_id=EMP_B,
                    claim_kind="insured",
                    product_code="GHS",
                    claim_type="outpatient",
                    incurred_date=_date(2026, 6, 1),
                    amount_claimed=100.0,
                    status="submitted",
                )
            )
            s.commit()


def test_claims_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/claims?policy_year_id={PY_B}")
    assert res.status_code == 404


def test_conversations_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    for suffix in ("", "&awaiting=any", f"&employee_id={EMP_B}"):
        res = client_as_a.get(
            f"/api/v1/conversations?policy_year_id={PY_B}{suffix}"
        )
        assert res.status_code == 404, suffix


def test_claim_get_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.get(f"/api/v1/claims/{CLAIM_B}")
    assert res.status_code == 404


def test_claim_decision_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(
        f"/api/v1/claims/{CLAIM_B}/decision", json={"action": "approve"}
    )
    assert res.status_code == 404


def test_claim_set_conversion_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(
        f"/api/v1/claims/{CLAIM_B}/conversion", json={"converted_amount": 100.0}
    )
    assert res.status_code == 404


def test_claim_fx_refresh_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(f"/api/v1/claims/{CLAIM_B}/fx-refresh")
    assert res.status_code == 404


def test_claim_send_to_insurer_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(f"/api/v1/claims/{CLAIM_B}/send-to-insurer", json={})
    assert res.status_code == 404


def test_claim_payment_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(
        f"/api/v1/claims/{CLAIM_B}/payment", json={"paid_on": "2030-01-01"}
    )
    assert res.status_code == 404


def test_claim_assessment_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.patch(
        f"/api/v1/claims/{CLAIM_B}/assessment", json={"taxable": True}
    )
    assert res.status_code == 404


def test_claim_amendment_cross_tenant_404(client_as_a: TestClient) -> None:
    """Correcting what the member stated is a separate handler from
    `/assessment`, and reaches further — it rewrites the claim's own figures.
    It goes through `load_claim` like the rest, so another tenant's claim is
    simply not found."""
    _ensure_claim_b()
    res = client_as_a.patch(
        f"/api/v1/claims/{CLAIM_B}", json={"amount_claimed": 1.0}
    )
    assert res.status_code == 404


def test_claim_document_download_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.get(f"/api/v1/claims/{CLAIM_B}/documents/any-doc/download")
    assert res.status_code == 404


def test_claim_review_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.get(f"/api/v1/claims/{CLAIM_B}/review")
    assert res.status_code == 404


def test_claim_rerun_review_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(f"/api/v1/claims/{CLAIM_B}/rerun-review")
    assert res.status_code == 404


def test_claim_messages_list_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.get(f"/api/v1/claims/{CLAIM_B}/messages")
    assert res.status_code == 404


def test_claim_message_post_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(
        f"/api/v1/claims/{CLAIM_B}/messages", json={"body": "cross-tenant"}
    )
    assert res.status_code == 404


def test_claim_messages_read_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(f"/api/v1/claims/{CLAIM_B}/messages/read")
    assert res.status_code == 404


def test_log_case_create_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/employees/{EMP_B}/log-cases",
        json={
            "claim_kind": "insured",
            "product_code": "GHS",
            "incurred_date": "2026-06-01",
            "amount_claimed": 100.0,
        },
    )
    assert res.status_code == 404


def test_claim_case_type_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.patch(
        f"/api/v1/claims/{CLAIM_B}/case-type",
        json={"case_type": "log", "reason": "cross-tenant"},
    )
    assert res.status_code == 404


def test_claim_document_upload_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.post(
        f"/api/v1/claims/{CLAIM_B}/documents",
        files={"file": ("x.pdf", b"%PDF-1.4 cross tenant", "application/pdf")},
    )
    assert res.status_code == 404


def test_enquiry_cross_tenant_404(client_as_a: TestClient) -> None:
    """A question belonging to client B, through client A's session. 404 on
    every verb — read, write, read-marker and status — so a broker cannot map
    another tenant's records, and cannot answer into them either."""
    from app.models import MemberEnquiry

    with SessionLocal() as s:
        if s.get(MemberEnquiry, ENQUIRY_B) is None:
            s.add(
                MemberEnquiry(
                    id=ENQUIRY_B,
                    client_id=CLIENT_B_ID,
                    employee_id=EMP_B,
                    policy_year_id=PY_B,
                    topic="coverage",
                    subject="B's question",
                    status="open",
                )
            )
            s.commit()

    assert client_as_a.get(f"/api/v1/enquiries/{ENQUIRY_B}").status_code == 404
    assert (
        client_as_a.get(f"/api/v1/enquiries/{ENQUIRY_B}/messages").status_code == 404
    )
    assert (
        client_as_a.post(
            f"/api/v1/enquiries/{ENQUIRY_B}/messages", json={"body": "hello"}
        ).status_code
        == 404
    )
    assert (
        client_as_a.post(
            f"/api/v1/enquiries/{ENQUIRY_B}/messages/read"
        ).status_code
        == 404
    )
    assert (
        client_as_a.post(
            f"/api/v1/enquiries/{ENQUIRY_B}/status", json={"action": "close"}
        ).status_code
        == 404
    )


def test_portal_preview_enquiry_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/employees/{EMP_B}/portal-preview/enquiries/{ENQUIRY_B}"
    )
    assert res.status_code == 404
    res = client_as_a.get(
        f"/api/v1/employees/{EMP_B}/portal-preview/enquiries/{ENQUIRY_B}/messages"
    )
    assert res.status_code == 404


def test_portal_preview_conversations_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/portal-preview/conversations")
    assert res.status_code == 404


def test_portal_preview_claim_messages_cross_tenant_404(
    client_as_a: TestClient,
) -> None:
    _ensure_claim_b()
    res = client_as_a.get(
        f"/api/v1/employees/{EMP_B}/portal-preview/claims/{CLAIM_B}/messages"
    )
    assert res.status_code == 404


def test_portal_preview_claim_cross_tenant_404(client_as_a: TestClient) -> None:
    _ensure_claim_b()
    res = client_as_a.get(
        f"/api/v1/employees/{EMP_B}/portal-preview/claims/{CLAIM_B}"
    )
    assert res.status_code == 404


def test_employee_utilization_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/utilization")
    assert res.status_code == 404


def test_dependant_approval_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/dependants/{DEP_B}/approval", json={"action": "approve"}
    )
    assert res.status_code == 404


def test_dependant_documents_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/dependants/{DEP_B}/documents")
    assert res.status_code == 404


# ── Panel clinic listings (clinic locator) ───────────────────────────────────
def test_panel_listing_clinics_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/panel-listings/{PANEL_B}/clinics")
    assert res.status_code == 404


def test_panel_listing_download_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/panel-listings/{PANEL_B}/download")
    assert res.status_code == 404


def test_panel_listing_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/panel-listings/{PANEL_B}", json={"label": "hijack"}
    )
    assert res.status_code == 404


def test_panel_listing_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/panel-listings/{PANEL_B}")
    assert res.status_code == 404


def test_panel_listing_upload_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/panel-listings/{PANEL_B}/upload",
        files={"file": ("panel.xlsx", b"stub", "application/octet-stream")},
    )
    assert res.status_code == 404


def test_panel_listing_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/panel-listings")
    assert res.status_code == 200
    assert PANEL_B not in {listing["id"] for listing in res.json()}


def test_policy_year_panels_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/panels")
    assert res.status_code == 404
    res = client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/panels", json={"panel_listing_ids": []}
    )
    assert res.status_code == 404


def test_policy_year_panels_cannot_tag_foreign_listing(client_as_a: TestClient) -> None:
    """Tagging client B's listing onto client A's own policy year must 404."""
    py_a = client_as_a.get("/api/v1/policy-years").json()[0]["id"]
    res = client_as_a.put(
        f"/api/v1/policy-years/{py_a}/panels",
        json={"panel_listing_ids": [PANEL_B]},
    )
    assert res.status_code == 404


def test_portal_preview_clinics_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/portal-preview/clinics")
    assert res.status_code == 404


def test_panel_card_cross_tenant_404(client_as_a: TestClient) -> None:
    """Card mutations go through load_panel_card — another tenant's PINNED
    card stays invisible (library cards, client_id NULL, are shared)."""
    assert client_as_a.patch(
        f"/api/v1/panel-cards/{CARD_B}", json={"name": "hijack"}
    ).status_code == 404
    assert client_as_a.delete(f"/api/v1/panel-cards/{CARD_B}").status_code == 404
    assert (
        client_as_a.get(f"/api/v1/panel-cards/{CARD_B}/artwork/front").status_code == 404
    )
    assert client_as_a.put(
        f"/api/v1/panel-cards/{CARD_B}/placements", json={"fields": []}
    ).status_code == 404


def test_panel_card_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    res = client_as_a.get("/api/v1/panel-cards")
    assert res.status_code == 200
    assert all(card["id"] != CARD_B for card in res.json())


def test_policy_year_cards_cross_tenant_404(client_as_a: TestClient) -> None:
    assert client_as_a.get(f"/api/v1/policy-years/{PY_B}/cards").status_code == 404
    assert client_as_a.post(
        f"/api/v1/policy-years/{PY_B}/cards",
        json={"panel_card_id": CARD_B, "product_id": PRODUCT_B},
    ).status_code == 404
    assert client_as_a.put(
        f"/api/v1/policy-years/{PY_B}/cards/{CARD_ASSIGNMENT_B}",
        json={"panel_card_id": CARD_B, "product_id": PRODUCT_B},
    ).status_code == 404
    assert client_as_a.delete(
        f"/api/v1/policy-years/{PY_B}/cards/{CARD_ASSIGNMENT_B}"
    ).status_code == 404


def test_card_assignment_probe_on_own_year_404s(client_as_a: TestClient) -> None:
    """An assignment id belonging to ANOTHER year, probed through a year the
    caller legitimately owns, must 404 (and is security-logged)."""
    py_a = client_as_a.get("/api/v1/policy-years").json()[0]["id"]
    assert client_as_a.delete(
        f"/api/v1/policy-years/{py_a}/cards/{CARD_ASSIGNMENT_B}"
    ).status_code == 404


def test_policy_year_cards_cannot_assign_foreign_card(client_as_a: TestClient) -> None:
    """Assigning client B's pinned card onto client A's own year must 404."""
    py_a = client_as_a.get("/api/v1/policy-years").json()[0]["id"]
    res = client_as_a.post(
        f"/api/v1/policy-years/{py_a}/cards",
        json={"panel_card_id": CARD_B, "product_id": PRODUCT_B},
    )
    assert res.status_code == 404


def test_portal_preview_cards_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/employees/{EMP_B}/portal-preview/cards")
    assert res.status_code == 404


def test_panel_listing_companies_cross_tenant_404(client_as_a: TestClient) -> None:
    """The multi-company enablement endpoints go through load_panel_listing —
    another tenant's PINNED listing stays invisible."""
    res = client_as_a.get(f"/api/v1/panel-listings/{PANEL_B}/companies")
    assert res.status_code == 404
    res = client_as_a.put(
        f"/api/v1/panel-listings/{PANEL_B}/companies", json={"client_ids": []}
    )
    assert res.status_code == 404


def test_insurers_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    """The catalog is tenant-or-global: client A sees the shared library but
    never client B's own entries."""
    rows = client_as_a.get("/api/v1/schemas/insurers").json()
    ids = {r["id"] for r in rows}
    assert INSURER_B not in ids
    # The seeded Singapore library IS visible (client_id NULL).
    assert any(r["client_id"] is None for r in rows)


def test_insurer_update_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/schemas/insurers/{INSURER_B}", json={"legal_name": "Hijacked"}
    )
    assert res.status_code == 404


def test_insurer_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/schemas/insurers/{INSURER_B}")
    assert res.status_code == 404


def test_entity_vocab_cross_tenant_404(client_as_a: TestClient) -> None:
    """The Insured picker's vocabulary exposes roster entity names — it must
    not leak another tenant's."""
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/entity-vocab")
    assert res.status_code == 404


def test_entity_alias_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    """Aliases are per-client — they name a client's own subsidiaries, so there
    is no shared library tier and another tenant's rows must be invisible."""
    rows = client_as_a.get("/api/v1/entity-aliases").json()
    assert all(r["id"] != ALIAS_B for r in rows)


def test_entity_alias_update_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/entity-aliases/{ALIAS_B}", json={"canonical": "Hijacked"}
    )
    assert res.status_code == 404


def test_entity_alias_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/entity-aliases/{ALIAS_B}")
    assert res.status_code == 404


def test_claim_doc_types_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    """The registry is per-client (lazy-seeded defaults) — client B's custom
    rows must never appear in client A's list."""
    rows = client_as_a.get("/api/v1/claim-doc-types").json()
    assert all(r["id"] != DOCTYPE_B for r in rows)
    # Client A got its own seeded defaults.
    assert {r["key"] for r in rows} >= {"discharge_summary", "finalised_tax_invoice"}


def test_claim_doc_type_update_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/claim-doc-types/{DOCTYPE_B}",
        json={"display": "Hijacked", "aliases": [], "key_fields": []},
    )
    assert res.status_code == 404


def test_claim_doc_type_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/claim-doc-types/{DOCTYPE_B}")
    assert res.status_code == 404


def test_claim_review_configs_list_excludes_other_tenant(client_as_a: TestClient) -> None:
    """The per-claim-type review rule setup is per-client — client B's rows
    must never appear in client A's list."""
    rows = client_as_a.get("/api/v1/claim-review-configs").json()
    assert all(r["id"] != REVIEW_CONFIG_B for r in rows)


def test_claim_review_config_update_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.put(
        f"/api/v1/claim-review-configs/{REVIEW_CONFIG_B}",
        json={
            "claim_kind": "insured",
            "claim_key": "GP",
            "display_label": "Hijacked",
            "field_maps": [
                {"portal_field": "amount_claimed", "document_field": "Total"}
            ],
        },
    )
    assert res.status_code == 404


def test_claim_review_config_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/claim-review-configs/{REVIEW_CONFIG_B}")
    assert res.status_code == 404


def test_claim_review_options_requires_a_tenant(
    client_as_no_tenant: TestClient,
) -> None:
    """The claim-type vocabulary is built from the ACTIVE client's benefit year
    — an unbound principal must not reach it."""
    assert (
        client_as_no_tenant.get("/api/v1/claim-review-configs/options").status_code
        == 400
    )


def test_claim_review_create_requires_a_tenant(
    client_as_no_tenant: TestClient,
) -> None:
    res = client_as_no_tenant.post(
        "/api/v1/claim-review-configs",
        json={
            "claim_kind": "insured",
            "claim_key": "GP",
            "display_label": "Unbound",
            "field_maps": [
                {"portal_field": "amount_claimed", "document_field": "Total"}
            ],
        },
    )
    assert res.status_code == 400


def test_claim_review_preview_requires_a_tenant(
    client_as_no_tenant: TestClient,
) -> None:
    res = client_as_no_tenant.post(
        "/api/v1/claim-review-configs/preview",
        json={
            "claim_kind": "insured",
            "claim_key": "GP",
            "display_label": "Unbound",
            "field_maps": [
                {"portal_field": "amount_claimed", "document_field": "Total"}
            ],
        },
    )
    assert res.status_code == 400


def test_claim_review_sources_exclude_active_and_other_firms(
    client_as_a: TestClient,
) -> None:
    """The import picker is server-authoritative — it must offer exactly what
    /import accepts: same broker firm, never the active company."""
    body = client_as_a.get("/api/v1/claim-review-configs/sources").json()
    ids = {c["id"] for c in body}
    assert DEMO_CLIENT_ID not in ids
    assert CLIENT_B_ID in ids  # same firm → importable
    client_b = next(company for company in body if company["id"] == CLIENT_B_ID)
    assert client_b["configured_count"] == 3  # unavailable BPROD row is excluded


def test_claim_review_import_source_other_firm_404(client_as_a: TestClient) -> None:
    """The import source must be accessible AND in the same broker firm — a
    rival firm's client 404s (on Postgres its rows live in another schema, so
    a cross-firm import would silently copy nothing)."""
    from app.models import BrokerFirm

    other_firm_id = "00000000-0000-0000-0000-0000000000f2"
    other_client_id = "00000000-0000-0000-0000-0000000000f3"
    with SessionLocal() as session:
        if session.get(BrokerFirm, other_firm_id) is None:
            session.add(BrokerFirm(id=other_firm_id, name="Rival Brokers Two"))
            session.add(
                Client(
                    id=other_client_id,
                    name="Rival-firm client two",
                    broker_firm_id=other_firm_id,
                )
            )
            session.commit()

    assert (
        client_as_a.get(f"/api/v1/claim-review-configs/from/{other_client_id}").status_code
        == 404
    )
    res = client_as_a.post(
        "/api/v1/claim-review-configs/import",
        json={"source_client_id": other_client_id, "config_ids": [REVIEW_CONFIG_B]},
    )
    assert res.status_code == 404


def test_claim_review_import_same_firm_source_allowed(client_as_a: TestClient) -> None:
    """Only shared current claim types are listed, with hospital scope context."""
    listed = client_as_a.get(f"/api/v1/claim-review-configs/from/{CLIENT_B_ID}").json()
    assert {r["id"] for r in listed} == {
        REVIEW_CONFIG_B,
        REVIEW_CONFIG_B_GOVT,
        REVIEW_CONFIG_B_PRIVATE,
    }
    assert REVIEW_CONFIG_B_UNAVAILABLE not in {r["id"] for r in listed}
    hospital_scopes = {
        row["scope_code"]: row["group_label"]
        for row in listed
        if row["scope_code"].startswith("ghs_hospitalisation_")
    }
    assert hospital_scopes == {
        "ghs_hospitalisation_govt": "Government hospital",
        "ghs_hospitalisation_private": "Private hospital",
    }
    res = client_as_a.post(
        "/api/v1/claim-review-configs/import",
        json={"source_client_id": CLIENT_B_ID, "config_ids": [REVIEW_CONFIG_B]},
    )
    assert res.status_code == 200
    imported = res.json()["imported"]
    assert len(imported) == 1
    new_id = imported[0]["id"]
    assert new_id != REVIEW_CONFIG_B  # a copy, not a shared row
    # Clean up so A's list-isolation test stays order-independent.
    assert client_as_a.delete(
        f"/api/v1/claim-review-configs/{new_id}",
        params={"expected_updated_at": imported[0]["updated_at"]},
    ).status_code == 204


def test_claim_review_import_rejects_product_missing_from_either_company(
    client_as_a: TestClient,
) -> None:
    """The list filter is repeated at write time so crafted IDs cannot bypass it."""
    res = client_as_a.post(
        "/api/v1/claim-review-configs/import",
        json={
            "source_client_id": CLIENT_B_ID,
            "config_ids": [REVIEW_CONFIG_B_UNAVAILABLE],
        },
    )
    assert res.status_code == 422
    assert res.json()["detail"]["code"] == "claim_type_not_available"


def test_dashboard_summary_excludes_other_firm(client_as_a: TestClient) -> None:
    """The firm Home roll-up is scoped to `accessible_clients` (the broker-firm
    boundary). A company in a DIFFERENT firm must never appear, even though its
    rows share the single SQLite schema in tests."""
    from app.models import BrokerFirm

    other_firm_id = "00000000-0000-0000-0000-0000000000f0"
    other_client_id = "00000000-0000-0000-0000-0000000000f1"
    with SessionLocal() as session:
        if session.get(BrokerFirm, other_firm_id) is None:
            session.add(BrokerFirm(id=other_firm_id, name="Rival Brokers"))
            session.add(
                Client(
                    id=other_client_id,
                    name="Rival-firm client",
                    broker_firm_id=other_firm_id,
                )
            )
            session.commit()

    body = client_as_a.get("/api/v1/dashboard/summary").json()
    ids = {c["id"] for c in body["companies"]}
    # A's own firm companies are present; the rival firm's client is not.
    assert DEMO_CLIENT_ID in ids
    assert CLIENT_B_ID in ids
    assert other_client_id not in ids
    assert body["firm"]["company_count"] == len(body["companies"])


# ── Slip / roster ingest write paths (cross-tenant policy_year_id) ──────────
def test_placement_slip_parse_cross_tenant_404(client_as_a: TestClient) -> None:
    # The tenant guard runs before the workbook is read, so a dummy file is fine.
    res = client_as_a.post(
        "/api/v1/placement-slips/parse",
        files={"file": ("slip.xlsx", b"not-a-real-workbook", "application/octet-stream")},
        data={"policy_year_id": PY_B},
    )
    assert res.status_code == 404


def test_employees_upload_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        "/api/v1/employees/upload",
        files={"file": ("roster.xlsx", b"x", "application/octet-stream")},
        data={"policy_year_id": PY_B},
    )
    assert res.status_code == 404


def test_dependants_upload_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        "/api/v1/dependants/upload",
        files={"file": ("deps.xlsx", b"x", "application/octet-stream")},
        data={"policy_year_id": PY_B},
    )
    assert res.status_code == 404


def test_dependants_auto_match_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.post(
        f"/api/v1/dependants/auto-match?policy_year_id={PY_B}"
    )
    assert res.status_code == 404


# ── Match override (cross-tenant employee_id) ───────────────────────────────
def test_match_override_cross_tenant_404(client_as_a: TestClient) -> None:
    # load_employee proves tenant ownership; EMP_B belongs to tenant B → 404.
    res = client_as_a.post(
        f"/api/v1/match-results/employees/{EMP_B}/override", json={}
    )
    assert res.status_code == 404


# ── Schemas CRUD (cross-tenant catalog rows → load_editable_global 404) ─────
def test_schemas_product_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/schemas/products/{PRODUCT_B_OWNED}",
        json={"display_name": "hijacked"},
    )
    assert res.status_code == 404


def test_schemas_product_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/schemas/products/{PRODUCT_B_OWNED}")
    assert res.status_code == 404


def test_schemas_attribute_patch_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.patch(
        f"/api/v1/schemas/employee-attributes/{ATTR_B}",
        json={"display_name": "hijacked"},
    )
    assert res.status_code == 404


def test_schemas_attribute_delete_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.delete(f"/api/v1/schemas/employee-attributes/{ATTR_B}")
    assert res.status_code == 404


def test_reports_bundle_list_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks")
    assert res.status_code == 404


def test_reports_bundle_download_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/workbooks/flex-wallet"
    )
    assert res.status_code == 404


def test_reports_wallet_utilisation_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/flex-wallet")
    assert res.status_code == 404


def test_reports_wallet_summary_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(
        f"/api/v1/policy-years/{PY_B}/reports/workbooks/flex-wallet"
    )
    assert res.status_code == 404


def test_reports_leaver_summary_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/leavers")
    assert res.status_code == 404


def test_reports_leaver_details_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/leavers")
    assert res.status_code == 404


def test_reports_insurance_claims_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/claims-register")
    assert res.status_code == 404


def test_reports_employee_claims_cross_tenant_404(client_as_a: TestClient) -> None:
    res = client_as_a.get(f"/api/v1/policy-years/{PY_B}/reports/workbooks/claims-register")
    assert res.status_code == 404
