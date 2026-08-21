"""Fact-Find ``.docx`` rendering — basis-of-cover table layout regressions."""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn

from app.services.fact_find_form import AGE_BANDS, BasisRow, BenefitLine, SectionContext
from app.services.fact_find_render import (
    _ctext,
    _fill_age_band,
    _fill_basis,
    _fill_family,
    _fill_highest_sum,
    _fill_presently_insured,
    _physical_cells,
)


def _merge_header(table) -> None:
    """Merge the first two header cells into one (the template's merged
    'Category of Employees / Designation' spanning the numbered + wide columns)."""
    table.rows[0].cells[0].merge(table.rows[0].cells[1])


def test_designation_lands_in_wide_column_not_numbered_cell() -> None:
    # Reproduces the bug where "SM to SVP" was written into the narrow numbered
    # cell and wrapped one character per line. It must land in the wide cell with
    # the numeral preserved beside it.
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    h = table.rows[0].cells
    h[0].text = "Category of Employees / Designation"
    h[2].text = "No. of Employees"
    h[3].text = "Basis of Cover"
    _merge_header(table)
    table.rows[1].cells[0].text = "i)"
    table.rows[2].cells[0].text = "ii)"

    sec = SectionContext(
        code="GTL",
        title="GTL",
        basis_rows=[
            BasisRow(designation="SM to SVP", num_employees=143, sum_insured="250,000"),
            BasisRow(designation="Manager", num_employees=100, sum_insured="60,000"),
        ],
    )
    _fill_basis(sec, table)

    assert _ctext(table.rows[1].cells[0]) == "i)"  # numeral kept, narrow
    assert _ctext(table.rows[1].cells[1]) == "SM to SVP"  # designation, wide
    assert _ctext(table.rows[1].cells[2]) == "143"
    assert _ctext(table.rows[1].cells[3]) == "250,000"
    assert _ctext(table.rows[2].cells[1]) == "Manager"


def _set_vmerge(cell, val: str | None) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    vm = tc_pr.makeelement(qn("w:vMerge"), {})
    if val:
        vm.set(qn("w:val"), val)
    tc_pr.append(vm)


def test_vertical_merge_writes_one_row_per_category() -> None:
    # The GCGP/GCSP page merges several physical rows per numbered category.
    # Each category must be written once (to the merge-origin row), not have its
    # value overwritten by later continuation rows.
    doc = Document()
    table = doc.add_table(rows=5, cols=4)  # header + two 2-row merged categories
    h = table.rows[0].cells
    h[0].text = "Category of Employees / Designation"
    h[2].text = "No. of Employees"
    h[3].text = "Basis of Cover"
    _merge_header(table)
    # Category i) spans rows 1-2, category ii) spans rows 3-4.
    table.rows[1].cells[0].text = "i)"
    table.rows[3].cells[0].text = "ii)"
    _set_vmerge(table.rows[1].cells[0], "restart")
    _set_vmerge(table.rows[2].cells[0], None)  # continuation
    _set_vmerge(table.rows[3].cells[0], "restart")
    _set_vmerge(table.rows[4].cells[0], None)  # continuation

    sec = SectionContext(
        code="GCGP_GCSP",
        title="GCGP",
        basis_rows=[
            BasisRow(designation="Managers", num_employees=50),
            BasisRow(designation="Officers", num_employees=30),
        ],
    )
    _fill_basis(sec, table)

    # Origin rows get the two distinct categories; continuation rows are skipped
    # (so the merged region is not overwritten with the next category).
    assert _ctext(table.rows[1].cells[1]) == "Managers"
    assert _ctext(table.rows[3].cells[1]) == "Officers"


def test_table_auto_expands_for_extra_categories() -> None:
    # The template ships 2 numbered rows but there are 5 categories; the table
    # must clone rows to fit them all and renumber i)…v).
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    h = table.rows[0].cells
    h[0].text = "Category of Employees / Designation"
    h[2].text = "No. of Employees"
    h[3].text = "Basis of Cover"
    _merge_header(table)
    table.rows[1].cells[0].text = "i)"
    table.rows[2].cells[0].text = "ii)"

    rows = [BasisRow(designation=f"Cat {n}", num_employees=n) for n in range(1, 6)]
    sec = SectionContext(code="GTL", title="GTL", basis_rows=rows)
    _fill_basis(sec, table)

    body = [r for r in table.rows[1:] if _ctext(r.cells[0])]
    assert len(body) == 5
    assert [_ctext(r.cells[0]) for r in body] == ["i)", "ii)", "iii)", "iv)", "v)"]
    assert [_ctext(r.cells[1]) for r in body] == [f"Cat {n}" for n in range(1, 6)]


def test_expanded_merged_block_clones_all_physical_rows() -> None:
    # When a category spans a 2-row vertical merge, expansion must clone the
    # whole block so a 3rd category gets its own merged region (not a stray row).
    doc = Document()
    table = doc.add_table(rows=5, cols=4)
    h = table.rows[0].cells
    h[0].text = "Category of Employees / Designation"
    h[2].text = "Plan Name"
    h[3].text = "Basis of Cover"
    _merge_header(table)
    table.rows[1].cells[0].text = "i)"
    table.rows[3].cells[0].text = "ii)"
    _set_vmerge(table.rows[1].cells[0], "restart")
    _set_vmerge(table.rows[2].cells[0], None)
    _set_vmerge(table.rows[3].cells[0], "restart")
    _set_vmerge(table.rows[4].cells[0], None)

    rows = [BasisRow(designation=f"Cat {n}", num_employees=n) for n in range(1, 4)]
    sec = SectionContext(code="GCGP_GCSP", title="GCGP", basis_rows=rows)
    _fill_basis(sec, table)

    # 1 header + 3 categories x 2 physical rows = 7 rows total.
    assert len(table.rows) == 7
    origins = [r for r in table.rows[1:] if _ctext(r.cells[1])]
    assert [_ctext(r.cells[1]) for r in origins] == ["Cat 1", "Cat 2", "Cat 3"]


def test_single_column_layout_renumbers_cloned_rows() -> None:
    # When the numeral and designation share one cell (un-merged header variant),
    # cloned rows must still get a fresh numeral, not the template slot's stale one.
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    h = table.rows[0].cells
    h[0].text = "Category of Employees / Designation"
    h[1].text = "No. of Employees"
    h[2].text = "Basis of Cover"
    table.rows[1].cells[0].text = "i)"  # single cell holds numeral + designation

    rows = [BasisRow(designation=f"Cat {n}", num_employees=n) for n in range(1, 4)]
    sec = SectionContext(code="GTL", title="GTL", basis_rows=rows)
    _fill_basis(sec, table)

    body = [_ctext(r.cells[0]) for r in table.rows[1:] if _ctext(r.cells[0])]
    assert body == ["i) Cat 1", "ii) Cat 2", "iii) Cat 3"]  # not "i) ... i) ... i) ..."


def test_clinical_benefits_come_from_plan_lines() -> None:
    doc = Document()
    table = doc.add_table(rows=4, cols=7)
    headers = table.rows[0].cells
    headers[0].text = "Category of Employees / Designation"
    headers[2].text = "Plan Name"
    headers[3].text = "Benefits provided"
    headers[4].text = "Panel"
    headers[5].text = "Max limit"
    headers[6].text = "Employee Co-insurance"
    _merge_header(table)
    table.rows[1].cells[0].text = "i)"
    _set_vmerge(table.rows[1].cells[0], "restart")
    _set_vmerge(table.rows[2].cells[0], None)
    _set_vmerge(table.rows[3].cells[0], None)

    lines = [
        BenefitLine("GPs", "Y", "50 per visit", "5"),
        BenefitLine("Specialist Care", "Y / N", "3,000"),
        BenefitLine("Diagnostic Tests", "Y / N", "Refer to Specialist Care"),
    ]
    sec = SectionContext(
        code="GCGP_GCSP",
        title="Clinical",
        basis_rows=[BasisRow("Managers", 50, "Plan 1", clinical_lines=lines)],
    )

    _fill_basis(sec, table)

    assert [_ctext(table.rows[i].cells[3]) for i in range(1, 4)] == [
        "GPs",
        "Specialist Care",
        "Diagnostic Tests",
    ]
    assert _ctext(table.rows[1].cells[2]) == "Plan 1"
    assert _ctext(table.rows[1].cells[5]) == "50 per visit"


def test_family_table_expands_from_available_product_plans() -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=5)
    table.rows[0].cells[0].text = "No. of Insured Members"
    for cell, text in zip(
        table.rows[1].cells,
        ("", "Employee Only", "Employee & Spouse", "Employee & Child(ren)", "Employee & Family"),
        strict=True,
    ):
        cell.text = text
    table.rows[2].cells[0].text = "Plan"
    sec = SectionContext(
        code="GHS",
        title="Hospital",
        available_plans=["Plan 1", "Plan 2", "Plan 10"],
        family_local={"Plan 2": {"EO": 2, "ES": 3, "EC": 4, "EF": 5}},
    )

    _fill_family(sec, table, sec.family_local)

    assert [_ctext(row.cells[0]) for row in table.rows[2:]] == [
        "Plan 1",
        "Plan 2",
        "Plan 10",
    ]
    assert [_ctext(cell) for cell in table.rows[3].cells[1:]] == ["2", "3", "4", "5"]


def test_presently_insured_does_not_overwrite_merged_labels() -> None:
    template = Document("app/templates/fact_find_form.docx")
    table = template.tables[4]
    sec = SectionContext(
        code="GTL",
        title="Term Life",
        insurer="Example Insurer",
        period_from="01/01/2026",
        period_to="31/12/2026",
    )

    _fill_presently_insured(sec, table)

    period_row = next(
        row for row in table.rows if _ctext(_physical_cells(row)[0]).startswith("Period")
    )
    values = [_ctext(cell) for cell in _physical_cells(period_row)]
    assert values == [
        "Period of Insurance:",
        "from",
        "01/01/2026",
        "to",
        "31/12/2026",
    ]


def test_count_only_age_table_does_not_overwrite_totals_with_sum_insured() -> None:
    template = Document("app/templates/fact_find_form.docx")
    table = template.tables[46]
    sec = SectionContext(
        code="GCGP_GCSP",
        title="Clinical",
        age_bands={AGE_BANDS[0][0]: (7, 30), AGE_BANDS[1][0]: (11, 48)},
    )

    _fill_age_band(sec, table)

    assert [_ctext(cell) for cell in _physical_cells(table.rows[-1])] == [
        "Total",
        "18",
        "78",
    ]


def test_highest_sum_uses_compact_amount_and_adjacent_range_box() -> None:
    template = Document("app/templates/fact_find_form.docx")
    table = template.tables[10]
    sec = SectionContext(
        code="GTL",
        title="Term Life",
        highest_sum_insured_age=45,
        highest_sum_insured=2_000_000,
    )

    _fill_highest_sum(sec, table)

    assert _ctext(_physical_cells(table.rows[1])[2]) == "Amt: S$2m"
    selected = _physical_cells(table.rows[9])
    assert _ctext(selected[-2]) == "X"
    assert _ctext(selected[0]) == ""
